"""GRS-0085 — C deliverable sections (ADR-0023).

The proposition heatmap (10 C modules, subject vs peers) and the differentiation-vs-rarity map
(Rare-present = asset, Common-absent = gap) render into the Platform Power Report when the run
carries a C result, and omit cleanly when it does not. Deliverables stay approval-gated; the methods
appendix documents the C provenance."""

from __future__ import annotations

import random
from datetime import UTC, date, datetime
from io import BytesIO
from uuid import uuid4

from bcap_contracts.assessments import (
    AssessmentDocument,
    SubcomponentRating,
    WidgetObservation,
)
from bcap_contracts.common import EvidenceGrade, MaturityLevel, NonScoreState
from bcap_contracts.deliverables import DeliverableMode
from bcap_contracts.predictions import CBenchmarkRow
from bcap_contracts.registry import load_registry
from docx import Document

from grassmarket.assessments.service import complete_c_subcomponents
from grassmarket.atlas import score
from grassmarket.atlas.draft_coefficients import draft_v1_coefficient_set
from grassmarket.atlas.montecarlo import draft_v1_uncertainty_model, run_monte_carlo
from grassmarket.deliverables.builder import DeliverableContext, build_platform_power_report
from tests._atlas_inputs import uniform_inputs

_REGISTRY = load_registry()
_MODEL = draft_v1_uncertainty_model()
_E3 = EvidenceGrade.E3_ARTIFACT
_NOW = datetime(2026, 7, 16, tzinfo=UTC)


def _paragraphs(data: bytes) -> str:
    doc = Document(BytesIO(data))
    text = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            text += "\n" + " | ".join(c.text for c in row.cells)
    return text


def _rare_and_common_widget():
    rare = next(w for w in _REGISTRY.c_widgets if w.rarity == "Rare")
    common = next(w for w in _REGISTRY.c_widgets if w.rarity == "Common")
    return rare, common


def _c_context(*, with_c: bool) -> DeliverableContext:
    bpl_coeffs = draft_v1_coefficient_set(_REGISTRY)
    bpl_inputs = uniform_inputs(_REGISTRY)
    unc = run_monte_carlo(bpl_inputs, bpl_coeffs, _REGISTRY, _MODEL, random.Random(1), draws=48)

    if not with_c:
        result = score(bpl_inputs, bpl_coeffs, _REGISTRY)
        return DeliverableContext(
            subject="Meridian",
            result=result,
            uncertainty=unc,
            coefficients=bpl_coeffs,
            uncertainty_version=_MODEL.version,
            generated_on=date(2026, 7, 16),
        )

    c_coeffs = draft_v1_coefficient_set(_REGISTRY, score_c=True)
    onboarding = _REGISTRY.require_c_module("CUST_ONBOARDING")
    c_doc = AssessmentDocument(
        c_subcomponents=tuple(
            SubcomponentRating(
                module_key=onboarding.key,
                subcomponent_key=s.key,
                level=MaturityLevel.ADVANCED,
                evidence_grade=_E3,
            )
            for s in onboarding.subcomponents  # rate the whole module so C is scoreable
        )
    )
    c_subs = complete_c_subcomponents(c_doc, _REGISTRY)
    c_inputs = bpl_inputs.model_copy(update={"c_subcomponents": c_subs})
    result = score(c_inputs, c_coeffs, _REGISTRY)

    rare, common = _rare_and_common_widget()
    peers = (
        CBenchmarkRow(
            id=uuid4(),
            peer_name="Saxo",
            profile_key="retail",
            c_index=0.42,
            module_scores={"CUST_ONBOARDING": 0.55},
            methodology_version="1.1",
            coefficient_version=c_coeffs.version,
            approved=True,
            approved_by=uuid4(),
            approved_at=_NOW,
            ingested_at=_NOW,
        ),
    )
    widgets = (
        WidgetObservation(widget_key=rare.key, present=True, ease=5, usability=4, depth=4),
        WidgetObservation(widget_key=common.key, present=False),
    )
    return DeliverableContext(
        subject="Meridian",
        result=result,
        uncertainty=unc,
        coefficients=c_coeffs,
        uncertainty_version=_MODEL.version,
        generated_on=date(2026, 7, 16),
        c_peers=peers,
        widgets=widgets,
        c_widget_defs=_REGISTRY.c_widgets,
    )


def test_c_sections_render_when_a_c_result_is_present() -> None:
    text = _paragraphs(
        build_platform_power_report(_c_context(with_c=True), DeliverableMode.DRAFT_INTERNAL)
    )
    assert "Customer Proposition (C)" in text
    assert "Differentiation vs. rarity" in text
    rare, common = _rare_and_common_widget()
    assert rare.name in text  # a Rare, present widget → differentiation asset
    assert common.name in text  # a Common, absent widget → table-stakes gap
    assert "Saxo" not in text  # peer NAMES are not printed; only the aggregate position is
    assert "ahead of 1" in text  # subject C beats the one peer at 0.42


def test_c_sections_are_omitted_cleanly_without_a_c_result() -> None:
    text = _paragraphs(
        build_platform_power_report(_c_context(with_c=False), DeliverableMode.DRAFT_INTERNAL)
    )
    assert "Customer Proposition (C)" not in text
    assert "Differentiation vs. rarity" not in text


def test_differentiation_map_classifies_rare_present_and_common_absent() -> None:
    text = _paragraphs(
        build_platform_power_report(_c_context(with_c=True), DeliverableMode.DRAFT_INTERNAL)
    )
    assert "Differentiation assets (Rare, present)" in text
    assert "Table-stakes gaps (Common, absent)" in text


def test_paywalled_widget_is_surfaced_separately() -> None:
    ctx = _c_context(with_c=True)
    rare, common = _rare_and_common_widget()
    paywalled = WidgetObservation(
        widget_key=common.key, present=False, state=NonScoreState.PRESENT_PAYWALLED
    )
    ctx = DeliverableContext(
        subject=ctx.subject,
        result=ctx.result,
        uncertainty=ctx.uncertainty,
        coefficients=ctx.coefficients,
        uncertainty_version=ctx.uncertainty_version,
        generated_on=ctx.generated_on,
        c_peers=ctx.c_peers,
        widgets=(WidgetObservation(widget_key=rare.key, present=True, ease=5), paywalled),
        c_widget_defs=ctx.c_widget_defs,
    )
    text = _paragraphs(build_platform_power_report(ctx, DeliverableMode.DRAFT_INTERNAL))
    assert "Present but gated / defective" in text
    assert "Present (Paywalled)" in text
