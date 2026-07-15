"""Rating Committee tests (GRS-0021, Methodology §8, ADR-0011).

Pins the exit criteria: a high-stakes rating without committee sign-off blocks BOTH finalisation and
a client pack; the approved rationale and any dissent render into the methods appendix; sign-off is
peer-only (a consultant cannot approve their own assessment; a non-committee principal cannot decide
at all). Plus unit coverage of the required-items derivation and the rating-matched gate.
"""

from __future__ import annotations

import random
from datetime import UTC, date, datetime
from io import BytesIO
from uuid import uuid4

from bcap_contracts.assessments import (
    AssessmentDocument,
    MetricEntry,
    PowerEntry,
    SubcomponentRating,
)
from bcap_contracts.committee import (
    CommitteeDecision,
    CommitteeDecisionStatus,
    CommitteeItem,
    CommitteeItemType,
)
from bcap_contracts.common import (
    EvidenceGrade,
    MaturityLevel,
    MetricConfidence,
    StrengthRating,
)
from bcap_contracts.registry import load_registry
from docx import Document

from grassmarket.assessments.service import compute_score
from grassmarket.atlas.committee import committee_blockers, required_committee_items
from grassmarket.atlas.draft_coefficients import draft_v1_coefficient_set
from grassmarket.atlas.montecarlo import draft_v1_uncertainty_model, elicited_v1_uncertainty_model
from grassmarket.deliverables.gate import CommitteePendingError
from grassmarket.deliverables.service import render_platform_power_report
from tests.committee_helpers import approve_committee_queue, committee_queue, seed_committee_member
from tests.conftest import SeededConsultant, auth_header
from tests.dual_rating_helpers import reach_consensus
from tests.test_assessment_lifecycle import _body, _scoreable_partial_doc

_REGISTRY = load_registry()
_MODEL = draft_v1_uncertainty_model()
# A client pack needs a client-usable uncertainty model too (GRS-0033 §7 gate); these tests isolate
# the COMMITTEE gate, so they pass a client-usable model to get past the uncertainty gate first.
_CLIENT_MODEL = elicited_v1_uncertainty_model()
_MODULE = "APP_SERVER"
_SUB = "APP_SERVER_SECURITY_COMPLIANCE"


def _scoreable(client, owner: SeededConsultant) -> str:
    aid = client.post("/assessments", json={}, headers=auth_header(owner)).json()["id"]
    client.put(
        f"/assessments/{aid}", json=_body(_scoreable_partial_doc()), headers=auth_header(owner)
    )
    return aid


def _score():
    return compute_score(
        _scoreable_partial_doc(),
        draft_v1_coefficient_set(_REGISTRY),
        _REGISTRY,
        _MODEL,
        random.Random(1),
    )


def _client_usable_set():
    return draft_v1_coefficient_set(_REGISTRY).model_copy(update={"client_usable": True})


_E3 = EvidenceGrade.E3_ARTIFACT


def _powers(strength: StrengthRating) -> tuple:
    return tuple(
        PowerEntry(
            power_key=p.key,
            benefit=strength,
            barrier=strength,
            benefit_grade=_E3,
            barrier_grade=_E3,
        )
        for p in _REGISTRY.powers
    )


_METRICS = (MetricEntry(metric_key="AUA", raw=1_000_000_000, confidence=MetricConfidence.AUDITED),)


def _established_power_doc() -> AssessmentDocument:
    """All seven powers Established/Established → strength Established (a high-stakes power)."""
    subs = (
        SubcomponentRating(
            module_key=_MODULE,
            subcomponent_key=_SUB,
            level=MaturityLevel.ADVANCED,
            evidence_grade=_E3,
        ),
    )
    return AssessmentDocument(
        subject="Established powers",
        subcomponents=subs,
        metrics=_METRICS,
        powers=_powers(StrengthRating.ESTABLISHED),
    )


def _frontier_module_doc() -> AssessmentDocument:
    """Every APP_SERVER subcomponent at Frontier/E3 → the module gate is Frontier (high-stakes)."""
    module = _REGISTRY.require_module(_MODULE)
    subs = tuple(
        SubcomponentRating(
            module_key=_MODULE,
            subcomponent_key=s.key,
            level=MaturityLevel.FRONTIER,
            evidence_grade=_E3,
        )
        for s in module.subcomponents
    )
    return AssessmentDocument(
        subject="Frontier module",
        subcomponents=subs,
        metrics=_METRICS,
        powers=_powers(StrengthRating.EMERGING),
    )


def _result_of(document: AssessmentDocument):
    return compute_score(
        document, draft_v1_coefficient_set(_REGISTRY), _REGISTRY, _MODEL, random.Random(1)
    ).result


def _decision(
    item: CommitteeItem,
    *,
    status: CommitteeDecisionStatus = CommitteeDecisionStatus.APPROVED,
    rationale: str = "Reviewed against the moat-duration rubric; the rating holds.",
    dissent_note: str | None = None,
) -> CommitteeDecision:
    now = datetime.now(UTC)
    return CommitteeDecision(
        id=uuid4(),
        owner_consultant_id=uuid4(),
        created_at=now,
        updated_at=now,
        assessment_id=uuid4(),
        item_type=item.item_type,
        item_key=item.item_key,
        rating=item.rating,
        status=status,
        rationale=rationale,
        dissent_note=dissent_note,
        decided_by_consultant_id=uuid4(),
        decided_at=now,
    )


def _docx_text(data: bytes) -> str:
    return "\n".join(p.text for p in Document(BytesIO(data)).paragraphs)


# --- Unit: required items + the rating-matched gate --------------------------------------


def test_required_items_flag_the_triad_above_none() -> None:
    items = required_committee_items(_score().result)
    triad = {i.item_key for i in items if i.item_type is CommitteeItemType.TRIAD}
    assert triad == {"economic_value", "perceived_value", "defence_value"}
    assert all(i.rating != "None" for i in items)


def test_gate_clears_only_when_every_item_has_a_matching_approval() -> None:
    items = required_committee_items(_score().result)
    assert committee_blockers(items, [])  # nothing approved → blocked
    approved = [_decision(i) for i in items]
    assert committee_blockers(items, approved) == []  # all approved → clear
    # A rejected item is not an approval — still blocked.
    rejected = [_decision(items[0], status=CommitteeDecisionStatus.REJECTED)] + approved[1:]
    assert committee_blockers(items, rejected)
    # An approval at a since-changed rating is stale — still blocked.
    stale = [_decision(i) for i in items]
    stale[0] = stale[0].model_copy(update={"rating": "Wide"})
    assert committee_blockers(items, stale)


# --- Finalisation gate -------------------------------------------------------------------


def test_committee_queue_lists_the_high_stakes_items(client, alice: SeededConsultant) -> None:
    aid = _scoreable(client, alice)
    entries = committee_queue(client, aid, alice)
    assert len(entries) == 3
    assert {e["item"]["item_type"] for e in entries} == {"triad"}
    assert all(e["decision"] is None for e in entries)  # all pending


def test_high_stakes_ratings_block_finalisation_without_signoff(
    client, alice: SeededConsultant
) -> None:
    aid = _scoreable(client, alice)
    reach_consensus(client, aid, alice, _MODULE, [(_SUB, MaturityLevel.ADVANCED)])
    # Consensus is done, but the triad ratings have no committee sign-off.
    resp = client.post(f"/assessments/{aid}/finalise", headers=auth_header(alice))
    assert resp.status_code == 409
    assert "Rating Committee sign-off incomplete" in resp.json()["detail"]


def test_a_rejected_high_stakes_item_still_blocks_finalisation(
    client, alice: SeededConsultant
) -> None:
    aid = _scoreable(client, alice)
    reach_consensus(client, aid, alice, _MODULE, [(_SUB, MaturityLevel.ADVANCED)])
    member = seed_committee_member(client)
    entries = committee_queue(client, aid, alice)
    # Approve two, REJECT one — finalisation is still refused.
    for entry in entries[:-1]:
        item = entry["item"]
        client.post(
            f"/assessments/{aid}/committee/decide",
            json={**item, "status": "approved", "rationale": "Holds."},
            headers=auth_header(member),
        )
    reject_item = entries[-1]["item"]
    client.post(
        f"/assessments/{aid}/committee/decide",
        json={**reject_item, "status": "rejected", "rationale": "Overstated; downgrade."},
        headers=auth_header(member),
    )
    resp = client.post(f"/assessments/{aid}/finalise", headers=auth_header(alice))
    assert resp.status_code == 409
    assert "Rating Committee sign-off incomplete" in resp.json()["detail"]


# --- Peer-only sign-off (scoping) --------------------------------------------------------


def test_a_committee_member_cannot_sign_off_their_own_assessment(client) -> None:
    member = seed_committee_member(client)
    aid = _scoreable(client, member)  # the member OWNS this one
    item = committee_queue(client, aid, member)[0]["item"]
    resp = client.post(
        f"/assessments/{aid}/committee/decide",
        json={**item, "status": "approved", "rationale": "Looks right to me."},
        headers=auth_header(member),
    )
    assert resp.status_code == 409
    assert "peer challenge" in resp.json()["detail"]


def test_committee_work_queue_lists_pending_and_is_members_only(
    client, alice: SeededConsultant
) -> None:
    """GRS-0061: a committee member finds work via GET /committee/queue — every in-progress
    assessment with pending high-stakes items. A plain consultant is refused (403)."""
    aid = _scoreable(client, alice)  # triad rates above None → high-stakes items pending
    member = seed_committee_member(client)

    q = client.get("/committee/queue", headers=auth_header(member))
    assert q.status_code == 200
    row = next((r for r in q.json() if r["assessment_id"] == aid), None)
    assert row is not None and row["pending_count"] >= 1

    # A non-committee consultant has no work-queue.
    assert client.get("/committee/queue", headers=auth_header(alice)).status_code == 403


def test_a_speculative_pre_approval_is_refused(client, alice: SeededConsultant) -> None:
    """GRS-0051: a member cannot pre-approve a rating the score has not reached — only an item that
    is currently required, at its current rating, may be decided. Otherwise the finalise gate could
    clear later at that rating with no contemporaneous review."""
    aid = _scoreable(client, alice)
    member = seed_committee_member(client)
    item = committee_queue(client, aid, alice)[0]["item"]

    # Same item, but at a rating it does not currently hold → speculative, refused.
    speculative = {**item, "rating": "Wide" if item["rating"] != "Wide" else "Established"}
    bad = client.post(
        f"/assessments/{aid}/committee/decide",
        json={
            **speculative,
            "status": "approved",
            "rationale": "Pre-approving ahead of the score.",
        },
        headers=auth_header(member),
    )
    assert bad.status_code == 409
    assert "currently-required" in bad.json()["detail"]

    # A bogus item key is likewise refused.
    ghost = client.post(
        f"/assessments/{aid}/committee/decide",
        json={**item, "item_key": "NOT_A_REAL_ITEM", "status": "approved", "rationale": "x."},
        headers=auth_header(member),
    )
    assert ghost.status_code == 409

    # The genuine (item, rating) still works.
    ok = client.post(
        f"/assessments/{aid}/committee/decide",
        json={
            **item,
            "status": "approved",
            "rationale": "Reviewed against the moat-duration rubric.",
        },
        headers=auth_header(member),
    )
    assert ok.status_code == 200


def test_a_non_committee_consultant_cannot_decide(
    client, alice: SeededConsultant, bob: SeededConsultant
) -> None:
    aid = _scoreable(client, alice)
    item = committee_queue(client, aid, alice)[0]["item"]
    resp = client.post(
        f"/assessments/{aid}/committee/decide",
        json={**item, "status": "approved", "rationale": "I say yes."},
        headers=auth_header(bob),  # a plain consultant, not committee
    )
    assert resp.status_code == 404


def test_a_stranger_cannot_view_the_committee_queue(
    client, alice: SeededConsultant, bob: SeededConsultant
) -> None:
    aid = _scoreable(client, alice)
    assert client.get(f"/assessments/{aid}/committee", headers=auth_header(bob)).status_code == 404


def test_a_committee_member_decision_appears_in_the_queue(client, alice: SeededConsultant) -> None:
    aid = _scoreable(client, alice)
    member = approve_committee_queue(client, aid, alice)
    entries = committee_queue(client, aid, alice)
    assert all(e["decision"] is not None for e in entries)
    assert all(e["decision"]["status"] == "approved" for e in entries)
    assert all(
        e["decision"]["decided_by_consultant_id"] == str(member.principal.consultant_id)
        for e in entries
    )


# --- Client-pack gate (deliverables) -----------------------------------------------------


def test_client_pack_is_refused_without_committee_signoff() -> None:
    """The client-pack path refuses independently of finalisation (defence-in-depth, ADR-0011)."""
    art = _score()
    try:
        render_platform_power_report(
            inputs=art.inputs,
            stored_result=art.result,
            coefficients=_client_usable_set(),
            registry=_REGISTRY,
            model=_CLIENT_MODEL,
            subject="Meridian",
            generated_on=date(2026, 7, 13),
            client_facing=True,
            committee_decisions=(),
        )
        raise AssertionError("expected CommitteePendingError")
    except CommitteePendingError as exc:
        assert "Rating Committee sign-off" in str(exc)


def test_client_pack_is_allowed_once_every_item_is_signed_off() -> None:
    art = _score()
    decisions = tuple(_decision(i) for i in required_committee_items(art.result))
    rendered = render_platform_power_report(
        inputs=art.inputs,
        stored_result=art.result,
        coefficients=_client_usable_set(),
        registry=_REGISTRY,
        model=_CLIENT_MODEL,
        subject="Meridian",
        generated_on=date(2026, 7, 13),
        client_facing=True,
        committee_decisions=decisions,
    )
    assert rendered.docx_bytes  # a real client pack was produced


def test_internal_draft_is_allowed_without_committee_signoff() -> None:
    art = _score()
    rendered = render_platform_power_report(
        inputs=art.inputs,
        stored_result=art.result,
        coefficients=draft_v1_coefficient_set(_REGISTRY),
        registry=_REGISTRY,
        model=_MODEL,
        subject="Meridian",
        generated_on=date(2026, 7, 13),
        client_facing=False,
        committee_decisions=(),
    )
    assert rendered.docx_bytes  # watermarked internal draft — no sign-off needed


# --- Rationale + dissent render into the appendix ----------------------------------------


def test_committee_rationale_and_dissent_render_into_the_appendix() -> None:
    art = _score()
    items = required_committee_items(art.result)
    econ = next(i for i in items if i.item_key == "economic_value")
    approved_rationale = "Durable cost advantage from custody scale; peer-challenged and upheld."
    dissent = "One member read the switching evidence as weaker; deferred to the majority."
    decisions = (
        _decision(econ, rationale=approved_rationale),
        _decision(
            next(i for i in items if i.item_key == "perceived_value"),
            rationale="Brand recognition supports the rating.",
            dissent_note=dissent,
        ),
    )
    rendered = render_platform_power_report(
        inputs=art.inputs,
        stored_result=art.result,
        coefficients=draft_v1_coefficient_set(_REGISTRY),
        registry=_REGISTRY,
        model=_MODEL,
        subject="Meridian",
        generated_on=date(2026, 7, 13),
        client_facing=False,
        committee_decisions=decisions,
    )
    text = _docx_text(rendered.docx_bytes)
    assert "Rating Committee decisions" in text  # the appendix section
    assert dissent in text  # dissent recorded (§8 audit evidence)
    assert approved_rationale in text  # the committee-approved triad rationale is the client text


# --- Power-Established+ and Module-Frontier gate branches (the exit criterion, by name) ---


def test_established_power_is_a_high_stakes_item_and_gated() -> None:
    result = _result_of(_established_power_doc())
    items = required_committee_items(result)
    powers = [i for i in items if i.item_type is CommitteeItemType.POWER]
    assert powers and all(i.rating in ("Established", "Wide") for i in powers)
    assert committee_blockers(powers, [])  # no sign-off → blocked
    assert committee_blockers(powers, [_decision(i) for i in powers]) == []  # signed off → clear


def test_frontier_module_is_a_high_stakes_item_and_gated() -> None:
    result = _result_of(_frontier_module_doc())
    items = required_committee_items(result)
    modules = [i for i in items if i.item_type is CommitteeItemType.MODULE]
    assert modules and all(i.rating == "Frontier" for i in modules)
    assert committee_blockers(modules, [])  # no sign-off → blocked
    assert committee_blockers(modules, [_decision(i) for i in modules]) == []


def test_client_pack_refused_on_a_frontier_module_without_signoff() -> None:
    """The client-pack gate bites on the MODULE branch too, not just the triad."""
    art = compute_score(
        _frontier_module_doc(),
        _client_usable_set(),
        _REGISTRY,
        _MODEL,
        random.Random(1),
    )
    try:
        render_platform_power_report(
            inputs=art.inputs,
            stored_result=art.result,
            coefficients=_client_usable_set(),
            registry=_REGISTRY,
            model=_CLIENT_MODEL,
            subject="Meridian",
            generated_on=date(2026, 7, 13),
            client_facing=True,
            committee_decisions=(),
        )
        raise AssertionError("expected CommitteePendingError")
    except CommitteePendingError as exc:
        assert "module" in str(exc).lower()


def test_an_established_power_assessment_finalises_only_after_signoff(
    client, alice: SeededConsultant
) -> None:
    aid = client.post("/assessments", json={}, headers=auth_header(alice)).json()["id"]
    client.put(
        f"/assessments/{aid}", json=_body(_established_power_doc()), headers=auth_header(alice)
    )
    reach_consensus(client, aid, alice, _MODULE, [(_SUB, MaturityLevel.ADVANCED)])
    # The queue now carries POWER items as well as the triad — all need sign-off.
    entries = committee_queue(client, aid, alice)
    assert any(e["item"]["item_type"] == "power" for e in entries)
    assert (
        client.post(f"/assessments/{aid}/finalise", headers=auth_header(alice)).status_code == 409
    )
    approve_committee_queue(client, aid, alice)
    assert (
        client.post(f"/assessments/{aid}/finalise", headers=auth_header(alice)).status_code == 200
    )


# --- More committee edge cases -----------------------------------------------------------


def test_a_decision_with_empty_rationale_is_422(client, alice: SeededConsultant) -> None:
    aid = _scoreable(client, alice)
    member = seed_committee_member(client)
    item = committee_queue(client, aid, alice)[0]["item"]
    resp = client.post(
        f"/assessments/{aid}/committee/decide",
        json={**item, "status": "approved", "rationale": ""},
        headers=auth_header(member),
    )
    assert resp.status_code == 422


def test_re_deciding_an_item_updates_it_in_place(client, alice: SeededConsultant) -> None:
    aid = _scoreable(client, alice)
    member = seed_committee_member(client)
    item = committee_queue(client, aid, alice)[0]["item"]
    for decision_status, rationale in (
        ("rejected", "Overstated."),
        ("approved", "On reflection, holds."),
    ):
        client.post(
            f"/assessments/{aid}/committee/decide",
            json={**item, "status": decision_status, "rationale": rationale},
            headers=auth_header(member),
        )
    # One row for the item, latest call wins — not a duplicate, not a constraint error.
    matching = [
        e
        for e in committee_queue(client, aid, alice)
        if e["item"]["item_key"] == item["item_key"] and e["decision"] is not None
    ]
    assert len(matching) == 1
    assert matching[0]["decision"]["status"] == "approved"
    assert matching[0]["decision"]["rationale"] == "On reflection, holds."


def test_committee_decisions_are_locked_after_finalisation(client, alice: SeededConsultant) -> None:
    aid = _scoreable(client, alice)
    reach_consensus(client, aid, alice, _MODULE, [(_SUB, MaturityLevel.ADVANCED)])
    member = approve_committee_queue(client, aid, alice)
    assert (
        client.post(f"/assessments/{aid}/finalise", headers=auth_header(alice)).status_code == 200
    )
    item = committee_queue(client, aid, alice)[0]["item"]
    resp = client.post(
        f"/assessments/{aid}/committee/decide",
        json={**item, "status": "rejected", "rationale": "Changed my mind."},
        headers=auth_header(member),
    )
    assert resp.status_code == 409
    assert "finalised" in resp.json()["detail"]


def test_a_committee_member_can_view_a_non_owned_queue(client, alice: SeededConsultant) -> None:
    aid = _scoreable(client, alice)
    member = seed_committee_member(client)
    resp = client.get(f"/assessments/{aid}/committee", headers=auth_header(member))
    assert resp.status_code == 200
    assert len(resp.json()) == 3  # a member (not the owner) sees the queue


def test_queue_is_empty_for_an_unscoreable_document(client, alice: SeededConsultant) -> None:
    aid = client.post("/assessments", json={}, headers=auth_header(alice)).json()["id"]  # empty
    assert client.get(f"/assessments/{aid}/committee", headers=auth_header(alice)).json() == []
