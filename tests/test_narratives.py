"""AI first-draft narrative tests (GRS-0017, PRD §5, ADR-0009).

'AI proposes, humans approve' as a runtime guarantee: proposals are drafted deterministically
(offline, no live calls), a human approves with an edit trail, a junior-tier author needs senior
sign-off (the quality gate), and a client-facing pack with any unapproved AI section is refused.
"""

from __future__ import annotations

import random
from datetime import UTC, datetime
from io import BytesIO
from uuid import uuid4

import pytest
from bcap_contracts.common import AssessorLevel, ConsultantTier, Role
from bcap_contracts.deliverables import DeliverableMode
from bcap_contracts.narratives import AINarrative, NarrativeSection, NarrativeStatus
from docx import Document
from pydantic import ValidationError

from grassmarket.assessments.service import compute_score
from grassmarket.atlas.draft_coefficients import draft_v1_coefficient_set
from grassmarket.atlas.montecarlo import draft_v1_uncertainty_model
from grassmarket.auth.security import create_access_token, hash_password
from grassmarket.data.repository import Repository
from grassmarket.deliverables.builder import AI_DRAFTED_LABEL, append_narrative_appendix
from grassmarket.deliverables.gate import (
    SeniorApprovalError,
    UnapprovedNarrativeError,
    assert_narratives_approved,
    assert_senior_approval,
)
from grassmarket.deliverables.narrative import (
    DRAFTER_VERSION,
    PROMPT_TEMPLATE_VERSION,
    TemplateNarrativeDrafter,
    context_from_result,
    edit_summary,
)
from tests.conftest import SeededConsultant, auth_header
from tests.test_deliverables import _doc

_REGISTRY_MODEL = draft_v1_uncertainty_model()


def _result():
    from bcap_contracts.registry import load_registry

    registry = load_registry()
    art = compute_score(
        _doc(graded=True),
        draft_v1_coefficient_set(registry),
        registry,
        _REGISTRY_MODEL,
        random.Random(20260713),
    )
    return art.result


def _narrative(section: NarrativeSection, *, approved: bool) -> AINarrative:
    when = datetime(2026, 7, 13, tzinfo=UTC)
    return AINarrative(
        id=uuid4(),
        owner_consultant_id=uuid4(),
        created_at=when,
        updated_at=when,
        deliverable_id=uuid4(),
        scoring_run_id=uuid4(),
        section=section,
        status=NarrativeStatus.APPROVED if approved else NarrativeStatus.PROPOSED,
        proposed_text="Proposed prose.",
        drafter_version=DRAFTER_VERSION,
        prompt_template_version=PROMPT_TEMPLATE_VERSION,
        author_tier=ConsultantTier.CONSULTANT,
        final_text="Final prose." if approved else None,
        approved_by_consultant_id=uuid4() if approved else None,
        approved_at=datetime(2026, 7, 13, 12, tzinfo=UTC) if approved else None,
        edit_summary="approved without edits" if approved else None,
    )


# ---------------------------------------------------------------- drafter (deterministic, offline)
def test_drafter_is_deterministic_and_versioned() -> None:
    ctx = context_from_result(_result(), "Meridian Securities")
    drafter = TemplateNarrativeDrafter()
    first = drafter.draft(NarrativeSection.INTERPRETATION, ctx)
    second = drafter.draft(NarrativeSection.INTERPRETATION, ctx)
    assert first == second  # same context → same prose, no live call
    assert "Meridian Securities" in first
    assert drafter.version == DRAFTER_VERSION
    assert drafter.prompt_template_version == PROMPT_TEMPLATE_VERSION


def test_drafter_covers_every_section() -> None:
    ctx = context_from_result(_result(), "Meridian")
    drafter = TemplateNarrativeDrafter()
    for section in NarrativeSection:
        assert drafter.draft(section, ctx).strip()


def test_edit_summary_reports_no_edits_and_edits() -> None:
    assert edit_summary("same", "same") == "approved without edits"
    diff = edit_summary("the draft line", "the edited line")
    assert "draft" in diff and "edited" in diff and diff != "approved without edits"


# ---------------------------------------------------------------- contract state machine
def test_approved_narrative_requires_full_trail() -> None:
    with pytest.raises(ValidationError):
        AINarrative(
            id=uuid4(),
            owner_consultant_id=uuid4(),
            created_at=datetime(2026, 7, 13, tzinfo=UTC),
            updated_at=datetime(2026, 7, 13, tzinfo=UTC),
            deliverable_id=uuid4(),
            scoring_run_id=uuid4(),
            section=NarrativeSection.COMMENTARY,
            status=NarrativeStatus.APPROVED,
            proposed_text="p",
            drafter_version="d",
            prompt_template_version="t",
            author_tier=ConsultantTier.CONSULTANT,
        )  # APPROVED with no approver/timestamp/final_text


def test_unapproved_narrative_forbids_approval_trail() -> None:
    with pytest.raises(ValidationError):
        AINarrative(
            id=uuid4(),
            owner_consultant_id=uuid4(),
            created_at=datetime(2026, 7, 13, tzinfo=UTC),
            updated_at=datetime(2026, 7, 13, tzinfo=UTC),
            deliverable_id=uuid4(),
            scoring_run_id=uuid4(),
            section=NarrativeSection.COMMENTARY,
            status=NarrativeStatus.PROPOSED,
            proposed_text="p",
            drafter_version="d",
            prompt_template_version="t",
            author_tier=ConsultantTier.CONSULTANT,
            approved_by_consultant_id=uuid4(),  # a proposal must not carry an approver
        )


# ---------------------------------------------------------------- the gate (unit)
def test_gate_refuses_client_pack_with_unapproved_narrative() -> None:
    narratives = [_narrative(NarrativeSection.INTERPRETATION, approved=False)]
    with pytest.raises(UnapprovedNarrativeError):
        assert_narratives_approved(narratives, client_facing=True)


def test_gate_allows_internal_draft_with_unapproved_narrative() -> None:
    narratives = [_narrative(NarrativeSection.INTERPRETATION, approved=False)]
    assert_narratives_approved(narratives, client_facing=False)  # internal is allowed (watermarked)


def test_gate_allows_client_pack_when_all_approved() -> None:
    narratives = [_narrative(s, approved=True) for s in NarrativeSection]
    assert_narratives_approved(narratives, client_facing=True)  # no raise


# ---------------------------------------------------------------- seniority (quality) gate
def test_senior_approval_gate() -> None:
    # A junior author self/peer-approving is refused; a Consultant-tier approver passes.
    with pytest.raises(SeniorApprovalError):
        assert_senior_approval(
            author_tier=ConsultantTier.VENTURE_ASSOCIATE,
            approver_tier=ConsultantTier.VENTURE_ASSOCIATE,
        )
    with pytest.raises(SeniorApprovalError):
        assert_senior_approval(
            author_tier=ConsultantTier.ADVISOR, approver_tier=ConsultantTier.ADVISOR
        )
    # Senior sign-off is fine, and a Consultant author may self-approve.
    assert_senior_approval(
        author_tier=ConsultantTier.VENTURE_ASSOCIATE, approver_tier=ConsultantTier.CONSULTANT
    )
    assert_senior_approval(
        author_tier=ConsultantTier.CONSULTANT, approver_tier=ConsultantTier.CONSULTANT
    )


# ---------------------------------------------------------------- builder (approval trail render)
def test_appendix_labels_ai_drafted_and_renders_approval_trail() -> None:
    doc = Document()
    approved = _narrative(NarrativeSection.INTERPRETATION, approved=True)
    pending = _narrative(NarrativeSection.RECOMMENDATION, approved=False)
    append_narrative_appendix(doc, [approved, pending], DeliverableMode.DRAFT_INTERNAL)
    text = "\n".join(p.text for p in doc.paragraphs)
    assert AI_DRAFTED_LABEL in text
    assert "Final prose." in text  # approved → final text
    assert "Approved by consultant" in text and "Edits: approved without edits" in text
    assert "Proposed prose." in text  # pending → proposal text
    assert "not client-usable until approved" in text


# ---------------------------------------------------------------- HTTP + persistence
def _deliverable_id(client, owner: SeededConsultant) -> str:
    from tests.test_deliverables import _engagement_with_finalised

    eid = _engagement_with_finalised(client, owner)
    return client.post(
        f"/engagements/{eid}/deliverables",
        json={"client_facing": False},
        headers=auth_header(owner),
    ).json()["id"]


@pytest.fixture
def senior(session_factory, settings) -> SeededConsultant:
    """A Consultant-tier consultant who may self-approve their own narratives (the senior path)."""
    from grassmarket.data.repository import Principal

    session = session_factory()
    try:
        stored = Repository(session).create_consultant(
            email="senior@bruntsfieldcapital.com",
            full_name="Senior",
            hashed_password=hash_password("correct-horse-battery-staple"),
            role=Role.CONSULTANT,
            tier=ConsultantTier.CONSULTANT,
            assessor_level=AssessorLevel.TRAINED,
        )
        session.commit()
    finally:
        session.close()
    token = create_access_token(
        settings,
        consultant_id=stored.id,
        email=stored.email,
        role=stored.role,
        tier=stored.tier,
        assessor_level=stored.assessor_level,
    )
    return SeededConsultant(
        stored=stored, principal=Principal(consultant_id=stored.id, role=stored.role), token=token
    )


@pytest.fixture
def senior_admin(session_factory, settings) -> SeededConsultant:
    """A Consultant-tier ADMIN — the governance reviewer who can approve a junior's narrative across
    ownership (ADR-0009 governance-visibility path)."""
    from grassmarket.data.repository import Principal

    session = session_factory()
    try:
        stored = Repository(session).create_consultant(
            email="chair@bruntsfieldcapital.com",
            full_name="Chair",
            hashed_password=hash_password("correct-horse-battery-staple"),
            role=Role.ADMIN,
            tier=ConsultantTier.CONSULTANT,
            assessor_level=AssessorLevel.TRAINED,
        )
        session.commit()
    finally:
        session.close()
    token = create_access_token(
        settings,
        consultant_id=stored.id,
        email=stored.email,
        role=stored.role,
        tier=stored.tier,
        assessor_level=stored.assessor_level,
    )
    return SeededConsultant(
        stored=stored, principal=Principal(consultant_id=stored.id, role=stored.role), token=token
    )


def test_http_senior_governance_reviewer_approves_junior_narrative(
    client, alice: SeededConsultant, senior_admin: SeededConsultant
) -> None:
    # alice is a Venture Associate; a Consultant-tier governance reviewer signs off her draft across
    # ownership — the reachable senior path (not a dead-end).
    did = _deliverable_id(client, alice)
    nid = client.post(
        f"/deliverables/{did}/narratives",
        json={"sections": ["recommendation"]},
        headers=auth_header(alice),
    ).json()[0]["id"]
    resp = client.post(f"/narratives/{nid}/approve", json={}, headers=auth_header(senior_admin))
    assert resp.status_code == 200
    body = resp.json()
    assert body["status"] == "approved"
    assert body["approved_by_consultant_id"] == str(senior_admin.stored.id)


def test_http_double_approval_is_conflict_not_500(client, senior: SeededConsultant) -> None:
    did = _deliverable_id(client, senior)
    nid = client.post(
        f"/deliverables/{did}/narratives",
        json={"sections": ["interpretation"]},
        headers=auth_header(senior),
    ).json()[0]["id"]
    assert (
        client.post(f"/narratives/{nid}/approve", json={}, headers=auth_header(senior)).status_code
        == 200
    )
    # A second approval is a conflict (409), never an unhandled 500.
    assert (
        client.post(f"/narratives/{nid}/approve", json={}, headers=auth_header(senior)).status_code
        == 409
    )


def test_http_empty_final_text_rejected(client, senior: SeededConsultant) -> None:
    did = _deliverable_id(client, senior)
    nid = client.post(
        f"/deliverables/{did}/narratives",
        json={"sections": ["commentary"]},
        headers=auth_header(senior),
    ).json()[0]["id"]
    resp = client.post(
        f"/narratives/{nid}/approve", json={"final_text": "   "}, headers=auth_header(senior)
    )
    assert resp.status_code == 422  # an approved section is never blank


def test_http_propose_drafts_all_sections(client, senior: SeededConsultant) -> None:
    did = _deliverable_id(client, senior)
    resp = client.post(f"/deliverables/{did}/narratives", json={}, headers=auth_header(senior))
    assert resp.status_code == 201
    body = resp.json()
    assert {n["section"] for n in body} == {s.value for s in NarrativeSection}
    assert all(n["status"] == "proposed" for n in body)
    assert all(n["drafter_version"] == DRAFTER_VERSION for n in body)


def test_http_senior_can_self_approve_with_trail(client, senior: SeededConsultant) -> None:
    did = _deliverable_id(client, senior)
    proposed = client.post(
        f"/deliverables/{did}/narratives",
        json={"sections": ["interpretation"]},
        headers=auth_header(senior),
    ).json()[0]
    resp = client.post(
        f"/narratives/{proposed['id']}/approve",
        json={"final_text": "Consultant-edited interpretation."},
        headers=auth_header(senior),
    )
    assert resp.status_code == 200
    approved = resp.json()
    assert approved["status"] == "approved"
    assert approved["approved_by_consultant_id"] == str(senior.stored.id)
    assert approved["final_text"] == "Consultant-edited interpretation."
    assert approved["approved_at"] is not None
    assert approved["edit_summary"] and approved["edit_summary"] != "approved without edits"


def test_http_junior_self_approval_refused(client, alice: SeededConsultant) -> None:
    # alice is a Venture Associate — she cannot sign off her own AI draft (the quality gate).
    did = _deliverable_id(client, alice)
    proposed = client.post(
        f"/deliverables/{did}/narratives",
        json={"sections": ["commentary"]},
        headers=auth_header(alice),
    ).json()[0]
    resp = client.post(f"/narratives/{proposed['id']}/approve", json={}, headers=auth_header(alice))
    assert resp.status_code == 409  # senior approval required


def test_http_client_download_refused_with_unapproved_narrative(
    client, senior: SeededConsultant
) -> None:
    did = _deliverable_id(client, senior)
    client.post(
        f"/deliverables/{did}/narratives",
        json={"sections": ["interpretation"]},
        headers=auth_header(senior),
    )
    # A client-facing download with an unapproved AI section is refused BEFORE rendering (409).
    resp = client.get(
        f"/deliverables/{did}/download?client_facing=true", headers=auth_header(senior)
    )
    assert resp.status_code == 409
    assert "not approved" in resp.json()["detail"]
    # The internal (non-client) download still works — unapproved drafts are allowed, watermarked.
    assert (
        client.get(f"/deliverables/{did}/download", headers=auth_header(senior)).status_code == 200
    )


def test_http_internal_download_renders_narrative_appendix(
    client, senior: SeededConsultant
) -> None:
    # The proposed narrative actually appears in the regenerated (internal) .docx, labelled.
    did = _deliverable_id(client, senior)
    client.post(
        f"/deliverables/{did}/narratives",
        json={"sections": ["interpretation"]},
        headers=auth_header(senior),
    )
    resp = client.get(f"/deliverables/{did}/download", headers=auth_header(senior))
    assert resp.status_code == 200
    doc = Document(BytesIO(resp.content))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert AI_DRAFTED_LABEL in text
    assert "AI-drafted narratives" in text
    assert "not client-usable until approved" in text  # proposed, not yet approved


def test_http_narratives_scoped(client, senior: SeededConsultant, bob: SeededConsultant) -> None:
    did = _deliverable_id(client, senior)
    nid = client.post(
        f"/deliverables/{did}/narratives",
        json={"sections": ["interpretation"]},
        headers=auth_header(senior),
    ).json()[0]["id"]
    # bob owns none of it → 404 on list, propose, approve.
    assert (
        client.get(f"/deliverables/{did}/narratives", headers=auth_header(bob)).status_code == 404
    )
    assert (
        client.post(
            f"/deliverables/{did}/narratives", json={}, headers=auth_header(bob)
        ).status_code
        == 404
    )
    assert (
        client.post(f"/narratives/{nid}/approve", json={}, headers=auth_header(bob)).status_code
        == 404
    )
    # The owner sees their own.
    assert (
        len(client.get(f"/deliverables/{did}/narratives", headers=auth_header(senior)).json()) == 1
    )
