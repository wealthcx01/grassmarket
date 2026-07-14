"""Diagnostic pack completion tests (GRS-0018, PRD §5).

All seven deliverable types generate from a finalised run: the five single-run documents added here
(Executive Summary, Infrastructure Heatmap, Technical Appendix, Workshop Output) plus Score
Evolution (multi-run). The load-bearing rule: Not Assessed cells in the heatmap are visually
distinct in the document XML — never red, never the same fill as Basic (Methodology §3.2).
"""

from __future__ import annotations

import random
from datetime import date
from io import BytesIO

import pytest
from bcap_contracts.assessments import SubcomponentRating
from bcap_contracts.common import EvidenceGrade, MaturityLevel, NonScoreState
from bcap_contracts.deliverables import DeliverableMode, DeliverableType
from bcap_contracts.registry import load_registry
from docx import Document

from grassmarket.assessments.service import compute_score
from grassmarket.atlas.draft_coefficients import draft_v1_coefficient_set
from grassmarket.atlas.montecarlo import draft_v1_uncertainty_model
from grassmarket.deliverables.builder import DeliverableContext
from grassmarket.deliverables.charts import evolution_lines, index_tornado, module_radar
from grassmarket.deliverables.evolution import EvolutionRun, build_score_evolution
from grassmarket.deliverables.heatmap import (
    _BAND_FILL,
    NOT_APPLICABLE_FILL,
    NOT_ASSESSED_FILL,
    build_infrastructure_heatmap,
    cell_fill,
)
from grassmarket.deliverables.reports import (
    build_executive_summary,
    build_technical_appendix,
    build_workshop_output,
)
from grassmarket.deliverables.service import render_diagnostic_document
from tests.conftest import SeededConsultant, auth_header
from tests.test_deliverables import _client_usable_set, _doc

_REGISTRY = load_registry()
_MODEL = draft_v1_uncertainty_model()
_DRAFT = DeliverableMode.DRAFT_INTERNAL


def _context(*, graded: bool = True) -> DeliverableContext:
    coeffs = draft_v1_coefficient_set(_REGISTRY)
    art = compute_score(_doc(graded=graded), coeffs, _REGISTRY, _MODEL, random.Random(20260713))
    return DeliverableContext(
        subject="Meridian Securities",
        result=art.result,
        uncertainty=art.uncertainty,
        coefficients=coeffs,
        uncertainty_version=_MODEL.version,
        generated_on=date(2026, 7, 13),
    )


def _all_states_context() -> DeliverableContext:
    """A context whose heatmap renders ALL FOUR states: a Basic cell, a higher band (Advanced, from
    `_doc`), a Not Applicable cell, and Not Assessed (the auto-filled remainder). So the §3.2
    distinctness rule is proven against RENDERED cells, not just module constants."""
    app = next(m for m in _REGISTRY.modules if m.key == "APP_SERVER")
    spare = [s.key for s in app.subcomponents if s.key != "APP_SERVER_SECURITY_COMPLIANCE"]
    base = _doc(graded=True)  # APP_SERVER_SECURITY_COMPLIANCE @ Advanced + rest Not Assessed
    extra = (
        SubcomponentRating(
            module_key="APP_SERVER",
            subcomponent_key=spare[0],
            level=MaturityLevel.BASIC,
            evidence_grade=EvidenceGrade.E3_ARTIFACT,
        ),
        SubcomponentRating(
            module_key="APP_SERVER",
            subcomponent_key=spare[1],
            state=NonScoreState.NOT_APPLICABLE,
        ),
    )
    doc = base.model_copy(update={"subcomponents": (*base.subcomponents, *extra)})
    coeffs = draft_v1_coefficient_set(_REGISTRY)
    art = compute_score(doc, coeffs, _REGISTRY, _MODEL, random.Random(1))
    return DeliverableContext(
        subject="Meridian",
        result=art.result,
        uncertainty=art.uncertainty,
        coefficients=coeffs,
        uncertainty_version=_MODEL.version,
        generated_on=date(2026, 7, 13),
    )


def _paras(data: bytes) -> str:
    doc = Document(BytesIO(data))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(c.text for c in row.cells)
    return "\n".join(parts)


# ----------------------------------------------------------- generation (every single-run type)
def test_each_single_run_document_generates() -> None:
    ctx = _context()
    for build in (
        build_executive_summary,
        build_infrastructure_heatmap,
        build_technical_appendix,
        build_workshop_output,
    ):
        data = build(ctx, _DRAFT)
        assert data[:2] == b"PK"


def test_executive_summary_is_board_ready() -> None:
    text = _paras(build_executive_summary(_context(), _DRAFT))
    assert "Executive Summary — Meridian Securities" in text
    assert "Platform value V" in text
    assert "Platform Power triad" in text


def test_technical_appendix_carries_versions_provenance_and_gates() -> None:
    text = _paras(build_technical_appendix(_context(), _DRAFT))
    assert "Coefficient set: v1-draft-pending-elicitation" in text
    assert "Weight provenance" in text and "review due" in text
    assert "Module gate reasoning" in text


def test_workshop_output_tolerates_partial_data() -> None:
    text = _paras(build_workshop_output(_context(graded=False), _DRAFT))
    assert "Pre-engagement mode" in text
    assert "to discuss:" in text  # Not Assessed subcomponents surfaced as discussion inputs
    assert "discussion prompts" in text.lower()


_RED_ALARM = {"FF0000", "C00000", "E00000", "FF3333"}


def _rating_fills(data: bytes) -> dict[str, set[str]]:
    doc = Document(BytesIO(data))
    fills_by_text: dict[str, set[str]] = {}
    for table in doc.tables:
        for row in table.rows:
            rating_cell = row.cells[1]  # the shaded Rating column
            fill = cell_fill(rating_cell)
            if fill is not None:
                fills_by_text.setdefault(rating_cell.text, set()).add(fill)
    return fills_by_text


# ----------------------------------------------------------- the heatmap XML rule (§3.2)
def test_all_four_states_are_pairwise_distinct_on_rendered_output() -> None:
    # The load-bearing §3.2 rule, proven against RENDERED cells (not just constants): a fixture with
    # all four states must render four PAIRWISE-DISTINCT fills, and none may be a red alarm.
    fills = _rating_fills(build_infrastructure_heatmap(_all_states_context(), _DRAFT))
    rendered = {}
    for label, expected in (
        ("Not Assessed", NOT_ASSESSED_FILL),
        ("Not Applicable", NOT_APPLICABLE_FILL),
        ("Basic", _BAND_FILL["Basic"]),
        ("Advanced", _BAND_FILL["Advanced"]),
    ):
        assert fills.get(label) == {expected}, f"{label} did not render its expected fill"
        rendered[label] = expected
    # Pairwise distinct on the actual rendered fills — Not Assessed is never Basic, never NA, etc.
    assert len(set(rendered.values())) == 4
    # And none of the four is a red alarm colour (§3.2: Not Assessed is not a failure).
    for fill in rendered.values():
        assert fill.upper() not in _RED_ALARM


def test_not_assessed_cells_are_visually_distinct() -> None:
    fills_by_text = _rating_fills(build_infrastructure_heatmap(_context(), _DRAFT))
    # Not Assessed cells exist and are all the distinct neutral grey.
    assert fills_by_text.get("Not Assessed") == {NOT_ASSESSED_FILL}
    # ...distinct from Basic, and never a red alarm colour.
    assert NOT_ASSESSED_FILL != _BAND_FILL["Basic"]
    assert NOT_ASSESSED_FILL.upper() not in _RED_ALARM
    # ...and distinct from every assessed band fill actually rendered.
    assessed_fills = {
        f
        for text, fs in fills_by_text.items()
        if text not in ("Not Assessed", "Not Applicable")
        for f in fs
    }
    assert NOT_ASSESSED_FILL not in assessed_fills


def test_heatmap_draft_is_watermarked() -> None:
    doc = Document(BytesIO(build_infrastructure_heatmap(_context(), _DRAFT)))
    header = "\n".join(p.text for s in doc.sections for p in s.header.paragraphs)
    assert "DRAFT — not client-usable" in header


# ----------------------------------------------------------- Score Evolution (multi-run)
def test_score_evolution_diffs_two_runs_with_version_annotation() -> None:
    ctx = _context()
    run1 = ctx.result
    bumped = run1.composite.model_copy(update={"v_index": round(run1.composite.v_index + 0.1, 6)})
    # Both methodology AND coefficient versions change between the runs — the annotation must name
    # each, so a movement is never misread as the client's own when it was a re-weighting.
    run2 = run1.model_copy(
        update={
            "composite": bumped,
            "coefficient_version": "v2-elicited",
            "methodology_version": "v1.3",
        }
    )
    data = build_score_evolution(
        subject="Meridian Securities",
        runs=[EvolutionRun("Baseline", run1), EvolutionRun("Re-score", run2)],
        mode=_DRAFT,
        generated_on=date(2026, 7, 13),
    )
    text = _paras(data)
    assert "Baseline" in text and "Re-score" in text
    assert "+10.0" in text  # ΔV of the bumped run (0.1 on the 0–100 display scale)
    assert "coefficients v1-draft-pending-elicitation→v2-elicited" in text  # version annotation
    assert f"methodology {run1.methodology_version}→v1.3" in text  # both changes named


def test_score_evolution_notes_absence_of_version_change() -> None:
    ctx = _context()
    text = _paras(
        build_score_evolution(
            subject="Meridian",
            runs=[EvolutionRun("R1", ctx.result), EvolutionRun("R2", ctx.result)],
            mode=_DRAFT,
            generated_on=date(2026, 7, 13),
        )
    )
    assert "No methodology or coefficient version change" in text


def test_score_evolution_needs_two_runs() -> None:
    ctx = _context()
    with pytest.raises(ValueError, match="at least two"):
        build_score_evolution(
            subject="Meridian",
            runs=[EvolutionRun("only", ctx.result)],
            mode=_DRAFT,
            generated_on=date(2026, 7, 13),
        )


# ----------------------------------------------------------- charts are deterministic
def test_charts_are_byte_deterministic() -> None:
    r1 = module_radar(labels=["A", "B", "C"], values=[40.0, 60.0, 80.0])
    r2 = module_radar(labels=["A", "B", "C"], values=[40.0, 60.0, 80.0])
    assert r1 == r2
    t1 = index_tornado(labels=["V"], lows=[40.0], mids=[50.0], highs=[60.0])
    t2 = index_tornado(labels=["V"], lows=[40.0], mids=[50.0], highs=[60.0])
    assert t1 == t2
    e1 = evolution_lines(run_labels=["a", "b"], series={"V": [1.0, 2.0]})
    e2 = evolution_lines(run_labels=["a", "b"], series={"V": [1.0, 2.0]})
    assert e1 == e2
    # The version stamp is actually stripped — no "matplotlib" tEXt chunk — so the byte-identity is
    # host-independent (within a pinned matplotlib), not just a same-process coincidence.
    assert b"matplotlib" not in r1
    assert b"matplotlib" not in t1


# ----------------------------------------------------------- service dispatcher + gate
def _render(dtype: DeliverableType, *, client_facing: bool, coeffs=None, model=None):
    from tests.committee_helpers import approved_decisions_for

    coeffs = coeffs or draft_v1_coefficient_set(_REGISTRY)
    model = model or _MODEL
    art = compute_score(_doc(graded=True), coeffs, _REGISTRY, model, random.Random(1))
    return render_diagnostic_document(
        deliverable_type=dtype,
        inputs=art.inputs,
        stored_result=art.result,
        coefficients=coeffs,
        registry=_REGISTRY,
        model=model,
        subject="Meridian",
        generated_on=date(2026, 7, 13),
        client_facing=client_facing,
        committee_decisions=approved_decisions_for(art.result),
    )


def test_dispatcher_renders_every_single_run_type() -> None:
    for dtype in (
        DeliverableType.EXECUTIVE_SUMMARY,
        DeliverableType.PLATFORM_POWER_REPORT,
        DeliverableType.INFRASTRUCTURE_HEATMAP,
        DeliverableType.TECHNICAL_APPENDIX,
        DeliverableType.WORKSHOP_OUTPUT,
    ):
        rendered = _render(dtype, client_facing=False)
        assert rendered.docx_bytes[:2] == b"PK"


def test_dispatcher_refuses_non_single_run_types() -> None:
    for dtype in (DeliverableType.MODERNISATION_ROADMAP, DeliverableType.SCORE_EVOLUTION):
        with pytest.raises(ValueError, match="own render path"):
            _render(dtype, client_facing=False)


def test_dispatcher_client_gate_still_applies() -> None:
    from grassmarket.atlas.montecarlo import elicited_v1_uncertainty_model
    from grassmarket.deliverables.gate import ClientUsabilityError

    with pytest.raises(ClientUsabilityError):
        _render(DeliverableType.EXECUTIVE_SUMMARY, client_facing=True)  # draft set
    # A client-usable coefficient set AND uncertainty model render a CLIENT executive summary.
    rendered = _render(
        DeliverableType.EXECUTIVE_SUMMARY,
        client_facing=True,
        coeffs=_client_usable_set(),
        model=elicited_v1_uncertainty_model(),
    )
    assert rendered.mode is DeliverableMode.CLIENT


def test_dispatcher_refuses_client_pack_on_draft_uncertainty_model() -> None:
    # The §7 twin of the coefficient gate (GRS-0033): client-usable coefficients are NOT enough —
    # a client pack's ranges must come from an elicited uncertainty model too. A draft model still
    # refuses even when the weights are client-usable.
    from grassmarket.atlas.montecarlo import elicited_v1_uncertainty_model
    from grassmarket.deliverables.gate import UncertaintyNotClientUsableError

    with pytest.raises(UncertaintyNotClientUsableError):
        _render(
            DeliverableType.TECHNICAL_APPENDIX,
            client_facing=True,
            coeffs=_client_usable_set(),
            model=draft_v1_uncertainty_model(),  # weights client-usable, widths draft → refuse
        )
    # Both client-usable → the client pack renders.
    rendered = _render(
        DeliverableType.TECHNICAL_APPENDIX,
        client_facing=True,
        coeffs=_client_usable_set(),
        model=elicited_v1_uncertainty_model(),
    )
    assert rendered.mode is DeliverableMode.CLIENT
    # A draft uncertainty model is fine for a watermarked internal draft.
    internal = _render(
        DeliverableType.TECHNICAL_APPENDIX,
        client_facing=False,
        model=draft_v1_uncertainty_model(),
    )
    assert internal.mode is DeliverableMode.DRAFT_INTERNAL


# ----------------------------------------------------------- HTTP (generate by type)
def _engagement(client, owner: SeededConsultant) -> str:
    from tests.test_deliverables import _engagement_with_finalised

    return _engagement_with_finalised(client, owner)


def test_http_generate_each_type_and_download(client, alice: SeededConsultant) -> None:
    eid = _engagement(client, alice)
    for value in (
        "executive_summary",
        "platform_power_report",
        "infrastructure_heatmap",
        "technical_appendix",
        "workshop_output",
    ):
        resp = client.post(
            f"/engagements/{eid}/deliverables",
            json={"client_facing": False, "deliverable_type": value},
            headers=auth_header(alice),
        )
        assert resp.status_code == 201, value
        assert resp.json()["type"] == value
        did = resp.json()["id"]
        dl = client.get(f"/deliverables/{did}/download", headers=auth_header(alice))
        assert dl.status_code == 200
        assert dl.content[:2] == b"PK"


def test_http_generate_non_single_run_types_refused(client, alice: SeededConsultant) -> None:
    eid = _engagement(client, alice)
    # Both types with their own render path are refused at generate (422), not 500 or silently.
    for value in ("modernisation_roadmap", "score_evolution"):
        resp = client.post(
            f"/engagements/{eid}/deliverables",
            json={"client_facing": False, "deliverable_type": value},
            headers=auth_header(alice),
        )
        assert resp.status_code == 422, value
