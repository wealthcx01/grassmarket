"""Deliverable core + client-usable gate tests (GRS-0015, PRD §5).

THE controlling gate: a client-facing pack may not be generated from a coefficient set with
``client_usable=False`` — a runtime refusal, tested. Draft coefficients emit only watermarked
"DRAFT — not client-usable" internal documents. Plus: honest uncertainty (an unmodelled B/P prints
a labelled point, never a false-tight band), the methods appendix carries versions + elicitation
dates, download regenerates the .docx, and every route is scoped.
"""

from __future__ import annotations

import random
from datetime import date
from io import BytesIO

import pytest
from bcap_contracts.assessments import (
    AssessmentDocument,
    CoefficientSet,
    MetricEntry,
    PowerEntry,
    SubcomponentRating,
)
from bcap_contracts.common import EvidenceGrade, MaturityLevel, MetricConfidence, StrengthRating
from bcap_contracts.deliverables import DeliverableMode
from bcap_contracts.registry import load_registry
from docx import Document

from grassmarket.assessments.service import compute_score
from grassmarket.atlas.draft_coefficients import draft_v1_coefficient_set
from grassmarket.atlas.montecarlo import draft_v1_uncertainty_model
from grassmarket.deliverables.builder import DeliverableContext, build_platform_power_report
from grassmarket.deliverables.gate import ClientUsabilityError, resolve_mode
from grassmarket.deliverables.service import render_platform_power_report
from tests.conftest import SeededConsultant, auth_header
from tests.test_engagement_detail import _contracted_prospect_http, _finalised_assessment_http

_REGISTRY = load_registry()
_MODEL = draft_v1_uncertainty_model()
_E3 = EvidenceGrade.E3_ARTIFACT


def _doc(*, graded: bool) -> AssessmentDocument:
    """A scoreable document. When ``graded`` is False, no metric confidence and no evidence grades —
    so B and P are UNMODELLED (their bands must print as labelled points, never a range)."""
    powers = tuple(
        PowerEntry(
            power_key=p.key,
            benefit=StrengthRating.EMERGING,
            barrier=StrengthRating.EMERGING,
            benefit_grade=_E3 if graded else None,
            barrier_grade=_E3 if graded else None,
        )
        for p in _REGISTRY.powers
    )
    metrics = (
        MetricEntry(
            metric_key="AUA",
            raw=1_000_000_000,
            confidence=MetricConfidence.AUDITED if graded else None,
        ),
    )
    # An assessed subcomponent always carries a grade (contract rule); B/P modelling depends only on
    # metric confidence + power grades, which is what `graded` toggles.
    subs = (
        SubcomponentRating(
            module_key="APP_SERVER",
            subcomponent_key="APP_SERVER_SECURITY_COMPLIANCE",
            level=MaturityLevel.ADVANCED,
            evidence_grade=_E3,
        ),
    )
    return AssessmentDocument(
        subject="Meridian", subcomponents=subs, metrics=metrics, powers=powers
    )


def _context(doc: AssessmentDocument, coefficients: CoefficientSet) -> DeliverableContext:
    art = compute_score(doc, coefficients, _REGISTRY, _MODEL, random.Random(20260706))
    return DeliverableContext(
        subject="Meridian Securities",
        result=art.result,
        uncertainty=art.uncertainty,
        coefficients=coefficients,
        uncertainty_version=_MODEL.version,
        generated_on=date(2026, 7, 9),
    )


def _client_usable_set() -> CoefficientSet:
    return draft_v1_coefficient_set(_REGISTRY).model_copy(update={"client_usable": True})


def test_non_retail_deliverable_renders_against_its_profile_view() -> None:
    # GRS-0148e regression: a wealth/exchange run's modules & metrics are profile-specific
    # (WEALTH_SUITABILITY, EXCH_ADV…). The deliverable must be built against that profile's registry
    # VIEW + coefficient set — building against the retail superset key-errors → a 500.
    from bcap_contracts.deliverables import DeliverableType

    from grassmarket.atlas import score
    from grassmarket.atlas.active import profile_scoring_context
    from grassmarket.deliverables.service import render_diagnostic_document
    from tests.test_atlas_engine_properties import build_inputs

    for profile in ("wealth", "exchange"):
        registry, coefficients = profile_scoring_context(profile)
        inputs = build_inputs(registry)
        result = score(inputs, coefficients, registry)
        for dtype in (DeliverableType.PLATFORM_POWER_REPORT, DeliverableType.EXECUTIVE_SUMMARY):
            rendered = render_diagnostic_document(
                deliverable_type=dtype,
                inputs=inputs,
                stored_result=result,
                coefficients=coefficients,
                registry=registry,
                model=_MODEL,
                subject="Segment Co",
                generated_on=date(2026, 7, 20),
                client_facing=False,
            )
            assert len(rendered.docx_bytes) > 1000  # a real .docx, not a crash


def _paragraphs(data: bytes) -> str:
    doc = Document(BytesIO(data))
    return "\n".join(p.text for p in doc.paragraphs)


# ------------------------------------------------------- the gate (runtime refusal)
def test_gate_refuses_client_pack_on_draft_set() -> None:
    draft = draft_v1_coefficient_set(_REGISTRY)
    assert draft.client_usable is False
    with pytest.raises(ClientUsabilityError):
        resolve_mode(draft, client_facing=True)


def test_gate_allows_internal_draft_on_draft_set() -> None:
    assert resolve_mode(draft_v1_coefficient_set(_REGISTRY), client_facing=False) is (
        DeliverableMode.DRAFT_INTERNAL
    )


def test_gate_allows_client_pack_on_client_usable_set() -> None:
    assert resolve_mode(_client_usable_set(), client_facing=True) is DeliverableMode.CLIENT


def test_service_refuses_client_facing_on_draft_set() -> None:
    ctx_doc = _doc(graded=True)
    art = compute_score(
        ctx_doc, draft_v1_coefficient_set(_REGISTRY), _REGISTRY, _MODEL, random.Random(1)
    )
    with pytest.raises(ClientUsabilityError):
        render_platform_power_report(
            inputs=art.inputs,
            stored_result=art.result,
            coefficients=draft_v1_coefficient_set(_REGISTRY),
            registry=_REGISTRY,
            model=_MODEL,
            subject="Meridian",
            generated_on=date(2026, 7, 9),
            client_facing=True,
        )


# ------------------------------------------------------- watermark + honest uncertainty
def test_draft_document_is_watermarked() -> None:
    data = build_platform_power_report(
        _context(_doc(graded=True), draft_v1_coefficient_set(_REGISTRY)),
        DeliverableMode.DRAFT_INTERNAL,
    )
    doc = Document(BytesIO(data))
    header_text = "\n".join(p.text for s in doc.sections for p in s.header.paragraphs)
    assert "DRAFT — not client-usable" in header_text


def test_modelled_indices_print_ranges() -> None:
    text = _paragraphs(
        build_platform_power_report(
            _context(_doc(graded=True), draft_v1_coefficient_set(_REGISTRY)),
            DeliverableMode.DRAFT_INTERNAL,
        )
    )
    assert "range" in text  # V/B/P/L all modelled → ranges


def test_unmodelled_bp_print_labelled_points_not_ranges() -> None:
    text = _paragraphs(
        build_platform_power_report(
            _context(_doc(graded=False), draft_v1_coefficient_set(_REGISTRY)),
            DeliverableMode.DRAFT_INTERNAL,
        )
    )
    # B and P had no confidence/grade → their statements are labelled points, not tight ranges.
    b_line = next(line for line in text.splitlines() if line.startswith("B (Business)"))
    p_line = next(line for line in text.splitlines() if line.startswith("P (Powers)"))
    assert "uncertainty not modelled" in b_line and "range" not in b_line
    assert "uncertainty not modelled" in p_line and "range" not in p_line
    # V is always modelled — it still prints a range.
    v_line = next(line for line in text.splitlines() if line.startswith("V (Platform Value)"))
    assert "range" in v_line


def test_methods_appendix_carries_versions_and_elicitation() -> None:
    text = _paragraphs(
        build_platform_power_report(
            _context(_doc(graded=True), draft_v1_coefficient_set(_REGISTRY)),
            DeliverableMode.DRAFT_INTERNAL,
        )
    )
    assert "Methods Appendix" in text
    assert "Coefficient set: v1-draft-pending-elicitation" in text
    # Under the DRAFT set the appendix must NOT claim elicitation — it states the honest placeholder
    # provenance (GRS-0033). "Expert-elicited" wording is reserved for the client-usable set.
    assert "Draft placeholder weights" in text and "not expert-elicited" in text
    assert "review due" in text


def test_methods_appendix_states_elicitation_under_client_usable_set() -> None:
    from grassmarket.atlas.elicited_coefficients import elicited_v1_coefficient_set

    text = _paragraphs(
        build_platform_power_report(
            _context(_doc(graded=True), elicited_v1_coefficient_set(_REGISTRY)),
            DeliverableMode.DRAFT_INTERNAL,
        )
    )
    assert "Coefficient set: v1-elicited-2026" in text
    assert "Weights expert-elicited" in text and "review due" in text


# ------------------------------------------------------- HTTP (generate / list / download / scope)
def _engagement_with_finalised(client, owner: SeededConsultant) -> str:
    pid = _contracted_prospect_http(client, owner)
    aid = _finalised_assessment_http(client, owner)
    return client.post(
        "/engagements",
        json={"prospect_id": pid, "title": "Delivery", "assessment_ids": [aid]},
        headers=auth_header(owner),
    ).json()["id"]


def test_http_generate_internal_draft(client, alice: SeededConsultant) -> None:
    eid = _engagement_with_finalised(client, alice)
    resp = client.post(
        f"/engagements/{eid}/deliverables",
        json={"client_facing": False},
        headers=auth_header(alice),
    )
    assert resp.status_code == 201
    body = resp.json()
    assert body["mode"] == "draft_internal"
    assert body["type"] == "platform_power_report"
    assert body["coefficient_version"] == "v1-draft-pending-elicitation"


def test_http_generate_client_facing_refused_on_draft(client, alice: SeededConsultant) -> None:
    eid = _engagement_with_finalised(client, alice)
    resp = client.post(
        f"/engagements/{eid}/deliverables", json={"client_facing": True}, headers=auth_header(alice)
    )
    assert resp.status_code == 409  # the gate — never a client pack on draft coefficients


def test_http_generate_without_finalised_assessment_refused(
    client, alice: SeededConsultant
) -> None:
    pid = _contracted_prospect_http(client, alice)
    eid = client.post(
        "/engagements", json={"prospect_id": pid, "title": "E"}, headers=auth_header(alice)
    ).json()["id"]
    resp = client.post(
        f"/engagements/{eid}/deliverables",
        json={"client_facing": False},
        headers=auth_header(alice),
    )
    assert resp.status_code == 409


def test_http_download_regenerates_docx(client, alice: SeededConsultant) -> None:
    eid = _engagement_with_finalised(client, alice)
    did = client.post(
        f"/engagements/{eid}/deliverables",
        json={"client_facing": False},
        headers=auth_header(alice),
    ).json()["id"]
    resp = client.get(f"/deliverables/{did}/download", headers=auth_header(alice))
    assert resp.status_code == 200
    assert "wordprocessingml" in resp.headers["content-type"]
    assert resp.content[:2] == b"PK"  # a real .docx (zip)


def test_http_deliverables_scoped(client, alice: SeededConsultant, bob: SeededConsultant) -> None:
    eid = _engagement_with_finalised(client, alice)
    did = client.post(
        f"/engagements/{eid}/deliverables",
        json={"client_facing": False},
        headers=auth_header(alice),
    ).json()["id"]

    assert (
        client.get(f"/engagements/{eid}/deliverables", headers=auth_header(bob)).status_code == 404
    )
    assert client.get(f"/deliverables/{did}/download", headers=auth_header(bob)).status_code == 404
    assert (
        client.post(
            f"/engagements/{eid}/deliverables",
            json={"client_facing": False},
            headers=auth_header(bob),
        ).status_code
        == 404
    )
    # Alice sees her own.
    assert (
        len(client.get(f"/engagements/{eid}/deliverables", headers=auth_header(alice)).json()) == 1
    )


# --- Assessment-level deliverable preview (GRS-0154) — no engagement needed --------------


def test_http_preview_finalised_assessment_returns_docx(client, alice: SeededConsultant) -> None:
    # The solo/sandbox "see the real deliverable" path: a finalised assessment previews its own
    # watermarked, internal-only deliverable without an engagement (mock-advisor: Priya/Elena).
    aid = _finalised_assessment_http(client, alice)
    url = f"/assessments/{aid}/deliverable-preview"
    resp = client.get(url, headers=auth_header(alice))
    assert resp.status_code == 200, resp.text
    assert "wordprocessingml" in resp.headers["content-type"]
    assert resp.content[:2] == b"PK"  # a real .docx (zip)
    # A preview is internal-only — the watermark must be present (never a client-facing pack).
    assert "DRAFT — not client-usable" in _paragraphs(resp.content)


def test_http_preview_unfinalised_assessment_refused(client, alice: SeededConsultant) -> None:
    aid = client.post("/assessments", json={"subject": "S"}, headers=auth_header(alice)).json()[
        "id"
    ]
    resp = client.get(f"/assessments/{aid}/deliverable-preview", headers=auth_header(alice))
    assert resp.status_code == 409  # finalise first — never render from an unfinalised run


def test_http_preview_is_owner_scoped(
    client, alice: SeededConsultant, bob: SeededConsultant
) -> None:
    aid = _finalised_assessment_http(client, alice)
    resp = client.get(f"/assessments/{aid}/deliverable-preview", headers=auth_header(bob))
    assert resp.status_code == 404  # another consultant cannot preview alice's assessment


# --- Deliverables index (GRS-0186) --------------------------------------------------------------

from datetime import UTC, datetime, timedelta  # noqa: E402

from bcap_contracts.deliverables import DeliverableType  # noqa: E402

from grassmarket.data.repository import Repository  # noqa: E402


def _engagement_with_deliverable(
    repo: Repository, owner, *, company: str, title: str, generated_at
):
    prospect = repo.create_prospect(owner.principal, company_name=company)
    from bcap_contracts.entities import PipelineStage

    for stage in (
        PipelineStage.WORKSHOP_SCHEDULED,
        PipelineStage.WORKSHOP_DELIVERED,
        PipelineStage.QUALIFIED,
        PipelineStage.SCOPED,
        PipelineStage.CONTRACTED,
    ):
        prospect = repo.update_prospect_stage(owner.principal, prospect.id, stage)
    eng = repo.create_engagement(owner.principal, prospect_id=prospect.id, title=title)
    deliverable = repo.create_deliverable(
        owner.principal,
        engagement_id=eng.id,
        deliverable_type=DeliverableType.EXECUTIVE_SUMMARY,
        title="Executive Summary",
        mode=DeliverableMode.DRAFT_INTERNAL,
        scoring_run_id=None,
        coefficient_version=None,
        content_hash=None,
        generated_at=generated_at,
    )
    return prospect, eng, deliverable


def test_deliverables_index_is_owner_scoped(repo: Repository, alice, bob) -> None:
    _engagement_with_deliverable(
        repo, alice, company="AliceCo", title="Alice engagement", generated_at=datetime.now(UTC)
    )
    _engagement_with_deliverable(
        repo, bob, company="BobCo", title="Bob engagement", generated_at=datetime.now(UTC)
    )
    alice_rows = repo.list_deliverables_for_consultant(alice.principal)
    bob_rows = repo.list_deliverables_for_consultant(bob.principal)
    assert [r.prospect_company_name for r in alice_rows] == ["AliceCo"]
    assert [r.prospect_company_name for r in bob_rows] == ["BobCo"]
    # Enriched fields resolve from the joined rows.
    row = alice_rows[0]
    assert row.engagement_title == "Alice engagement"
    assert row.prospect_id is not None
    assert row.type == DeliverableType.EXECUTIVE_SUMMARY


def test_deliverables_index_orders_newest_first_nulls_last(repo: Repository, alice) -> None:
    now = datetime.now(UTC)
    _engagement_with_deliverable(
        repo, alice, company="Older", title="e1", generated_at=now - timedelta(days=2)
    )
    _engagement_with_deliverable(repo, alice, company="Newer", title="e2", generated_at=now)
    _engagement_with_deliverable(repo, alice, company="Ungenerated", title="e3", generated_at=None)
    order = [
        r.prospect_company_name for r in repo.list_deliverables_for_consultant(alice.principal)
    ]
    assert order == ["Newer", "Older", "Ungenerated"]  # newest generated first, null last


def test_http_deliverables_index_is_self_scoped(client, alice, bob) -> None:
    from tests.conftest import auth_header

    # Seed via the repo behind a fresh session bound to the same engine.
    session = client.app.state.session_factory()
    try:
        r = Repository(session)
        _engagement_with_deliverable(
            r, alice, company="AliceCo", title="A", generated_at=datetime.now(UTC)
        )
        _engagement_with_deliverable(
            r, bob, company="BobCo", title="B", generated_at=datetime.now(UTC)
        )
        session.commit()
    finally:
        session.close()

    a = client.get("/deliverables", headers=auth_header(alice))
    assert a.status_code == 200
    assert [row["prospect_company_name"] for row in a.json()] == ["AliceCo"]
    # Cross-advisor negative: alice never sees bob's row.
    assert all(row["prospect_company_name"] != "BobCo" for row in a.json())


def test_http_deliverables_index_empty_for_new_advisor(client, alice) -> None:
    from tests.conftest import auth_header

    resp = client.get("/deliverables", headers=auth_header(alice))
    assert resp.status_code == 200 and resp.json() == []


def test_portfolio_row_linked_prospect_id(client, alice) -> None:
    """The portfolio row links to the client record only when the assessment is linked to an
    engagement (GRS-0186); an unlinked assessment stays None."""
    from tests.conftest import auth_header
    from tests.test_engagement_detail import (
        _contracted_prospect_http,
        _finalised_assessment_http,
    )

    headers = auth_header(alice)
    # An unlinked assessment (no engagement references it).
    unlinked_id = client.post(
        "/assessments", json={"subject": "Unlinked Co"}, headers=headers
    ).json()["id"]
    # A linked one: finalise it, then open an engagement referencing it.
    pid = _contracted_prospect_http(client, alice)
    linked_id = _finalised_assessment_http(client, alice)
    eng = client.post(
        "/engagements",
        json={"prospect_id": pid, "title": "Linked engagement", "assessment_ids": [linked_id]},
        headers=headers,
    ).json()
    rows = {
        r["assessment_id"]: r for r in client.get("/assessments/portfolio", headers=headers).json()
    }
    assert rows[unlinked_id]["linked_prospect_id"] is None
    assert rows[linked_id]["linked_prospect_id"] == eng["prospect_id"]
