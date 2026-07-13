"""Modernisation Roadmap tests (GRS-0016, §10) — the flagship money pages, honest by construction.

Generates the Roadmap from the golden-master run plus a worked scenario fixture and proves: the
value bridge renders with every lever NPV traceable to a printed assumption; the priority-vs-cost
scatter is embedded; the scenario comparison table carries every named scenario; priority (ΔV) and
price (currency) sit side by side but are never divided into one ROI number; and the client-usable
gate + watermark are inherited from GRS-0015 (a draft set → watermarked internal only).
"""

from __future__ import annotations

import random
from datetime import date
from io import BytesIO

import pytest
from bcap_contracts.common import MaturityLevel, StrengthRating
from bcap_contracts.deliverables import DeliverableMode
from bcap_contracts.money import Currency, Money
from bcap_contracts.registry import load_registry
from bcap_contracts.value import (
    Assumption,
    AssumptionRegister,
    ValueBridge,
)
from docx import Document

from grassmarket.assessments.service import compute_score
from grassmarket.atlas.draft_coefficients import draft_v1_coefficient_set
from grassmarket.atlas.montecarlo import draft_v1_uncertainty_model
from grassmarket.deliverables.gate import ClientUsabilityError
from grassmarket.deliverables.money_text import format_money
from grassmarket.deliverables.roadmap import (
    RoadmapContext,
    RoadmapEntry,
    build_modernisation_roadmap,
)
from grassmarket.deliverables.service import render_modernisation_roadmap
from grassmarket.value import (
    cost_estimate,
    cost_to_serve_lever,
    evaluate_scenario,
    incident_expected_loss_lever,
    prioritise_upgrades,
    project_drag_lever,
    revenue_enablement_lever,
    strategic_rating,
)
from tests._atlas_inputs import meridian_inputs, override
from tests.test_deliverables import _client_usable_set, _doc

_REGISTRY = load_registry()
_MODEL = draft_v1_uncertainty_model()
_GBP = Currency.GBP
_L = MaturityLevel


def _register() -> AssumptionRegister:
    return AssumptionRegister(
        entries=(
            Assumption(
                ref="AR-EFFORT",
                description="Remediation effort",
                baseline_value=180,
                unit="person-days",
                source="Delivery estimate, Meridian CTO workshop",
            ),
            Assumption(
                ref="AR-RATE",
                description="Blended day rate",
                baseline_value=900,
                unit="GBP/day",
                source="Meridian vendor rate card FY26",
            ),
            Assumption(
                ref="AR-CTS",
                description="Cost to serve per active client",
                baseline_value=140,
                unit="GBP/yr",
                source="Meridian management accounts FY25",
            ),
            Assumption(
                ref="AR-CLIENTS",
                description="Active clients",
                baseline_value=180_000,
                unit="count",
                source="Meridian KPI pack FY25",
            ),
            Assumption(
                ref="AR-DELIVERY",
                description="Annual delivery spend",
                baseline_value=12_000_000,
                unit="GBP/yr",
                source="Meridian technology budget FY26",
            ),
            Assumption(
                ref="AR-INCIDENT",
                description="Annual incident expected loss",
                baseline_value=2_500_000,
                unit="GBP/yr",
                source="Meridian risk register FY25",
            ),
            Assumption(
                ref="AR-REVENUE",
                description="Revenue enabled by faster settlement",
                baseline_value=6_000_000,
                unit="GBP/yr",
                source="Meridian commercial model (client-supplied)",
            ),
            Assumption(
                ref="AR-COST-OEMS",
                description="Indicative cost — fix OEMS execution algos",
                baseline_value=162_000,
                unit="GBP",
                source="Delivery estimate, Meridian CTO workshop",
            ),
            Assumption(
                ref="AR-COST-OBS",
                description="Indicative cost — top-up observability",
                baseline_value=40_000,
                unit="GBP",
                source="Delivery estimate, Meridian CTO workshop",
            ),
        )
    )


def _bridge() -> ValueBridge:
    reg = _register()
    cost = cost_estimate(
        effort_days=180,
        day_rate=900,
        currency=_GBP,
        primary_ref="AR-EFFORT",
        assumption_refs=["AR-EFFORT", "AR-RATE"],
        note="Back-office remediation.",
    )
    levers = (
        cost_to_serve_lever(
            cost_to_serve_per_client=140,
            active_clients=180_000,
            reduction_fraction=0.08,
            discount_rate=0.10,
            horizon_years=5,
            adoption_probability=0.7,
            currency=_GBP,
            primary_ref="AR-CTS",
            assumption_refs=["AR-CTS", "AR-CLIENTS"],
        ),
        project_drag_lever(
            annual_delivery_spend=12_000_000,
            drag_fraction=0.15,
            reduction_fraction=0.25,
            discount_rate=0.10,
            horizon_years=5,
            adoption_probability=0.6,
            currency=_GBP,
            primary_ref="AR-DELIVERY",
            assumption_refs=["AR-DELIVERY"],
        ),
        incident_expected_loss_lever(
            annual_incident_expected_loss=2_500_000,
            reduction_fraction=0.40,
            discount_rate=0.10,
            horizon_years=5,
            adoption_probability=0.5,
            currency=_GBP,
            primary_ref="AR-INCIDENT",
            assumption_refs=["AR-INCIDENT"],
        ),
        revenue_enablement_lever(
            enabled_annual_revenue=6_000_000,
            contribution_margin=0.55,
            discount_rate=0.10,
            horizon_years=5,
            realisation_probability=0.4,
            currency=_GBP,
            primary_ref="AR-REVENUE",
            assumption_refs=["AR-REVENUE"],
        ),
    )
    strategic = (
        strategic_rating(
            "Moat durability",
            StrengthRating.ESTABLISHED,
            "Custody + settlement reliability deepens switching costs over 3–5 years.",
        ),
    )
    return ValueBridge(
        subject="Meridian Securities — Back Office remediation",
        assumption_register=reg,
        cost=cost,
        levers=levers,
        strategic=strategic,
    )


def _scenario_defs():
    baseline = meridian_inputs()
    return baseline, [
        ("Fix OEMS execution algos", override(baseline, "OEMS_EXEC_ALGOS", _L.ADVANCED)),
        ("Top-up observability", override(baseline, "APP_SERVER_OBSERVABILITY", _L.FRONTIER)),
    ]


def _context() -> RoadmapContext:
    coeffs = draft_v1_coefficient_set(_REGISTRY)
    baseline, defs = _scenario_defs()
    priority = prioritise_upgrades(baseline, defs, coeffs, _REGISTRY)
    scenarios = tuple(
        evaluate_scenario(name, baseline, scen, coeffs, _REGISTRY) for name, scen in defs
    )
    # Each ranked intervention gets a register-linked indicative cost (Money) for the ranked table
    # and the priority-vs-cost scatter — keyed by intervention name so the ref matches the register.
    costs = {
        "Fix OEMS execution algos": Money(
            amount_minor=162_000_00, currency=_GBP, assumption_register_ref="AR-COST-OEMS"
        ),
        "Top-up observability": Money(
            amount_minor=40_000_00, currency=_GBP, assumption_register_ref="AR-COST-OBS"
        ),
    }
    entries = tuple(
        RoadmapEntry(name=u.name, rank=u.rank, delta_v=u.delta_v, cost=costs[u.name])
        for u in priority
    )
    art = compute_score(_doc(graded=True), coeffs, _REGISTRY, _MODEL, random.Random(20260706))
    return RoadmapContext(
        subject="Meridian Securities",
        bridge=_bridge(),
        entries=entries,
        scenarios=scenarios,
        engine_version=art.result.engine_version,
        methodology_version=art.result.methodology_version,
        coefficient_version=art.result.coefficient_version,
        uncertainty_version=_MODEL.version,
        generated_on=date(2026, 7, 13),
    )


def _all_text(data: bytes) -> str:
    """All paragraph text AND table-cell text (the roadmap puts most content in tables)."""
    doc = Document(BytesIO(data))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


# ---------------------------------------------------------------- generation from the golden run
def test_roadmap_generates_a_real_docx() -> None:
    data = build_modernisation_roadmap(_context(), DeliverableMode.DRAFT_INTERNAL)
    assert data[:2] == b"PK"  # a real .docx (zip)
    assert "Modernisation Roadmap — Meridian Securities" in _all_text(data)


def test_every_lever_npv_traces_to_a_printed_assumption() -> None:
    ctx = _context()
    text = _all_text(build_modernisation_roadmap(ctx, DeliverableMode.DRAFT_INTERNAL))
    register_refs = ctx.bridge.assumption_register.refs()
    for lever in ctx.bridge.levers:
        for ref in lever.assumption_refs:
            assert ref in register_refs  # traceability holds in the model
            assert ref in text  # ...and the register row is actually printed
    # Every register entry (ref + its source) renders.
    for a in ctx.bridge.assumption_register.entries:
        assert a.ref in text and a.source in text


def test_scatter_chart_is_embedded() -> None:
    data = build_modernisation_roadmap(_context(), DeliverableMode.DRAFT_INTERNAL)
    doc = Document(BytesIO(data))
    assert len(doc.inline_shapes) >= 1  # the priority-vs-cost scatter PNG


def test_scenario_comparison_lists_every_named_scenario() -> None:
    ctx = _context()
    text = _all_text(build_modernisation_roadmap(ctx, DeliverableMode.DRAFT_INTERNAL))
    for scenario in ctx.scenarios:
        assert scenario.name in text
    assert "ΔV" in text and "ΔL" in text


def test_priority_and_price_shown_side_by_side_never_one_ratio() -> None:
    data = build_modernisation_roadmap(_context(), DeliverableMode.DRAFT_INTERNAL)
    doc = Document(BytesIO(data))
    # Structural, not prose: the ranked table has ΔV and cost as DISTINCT columns.
    ranked = next(t for t in doc.tables if t.rows[0].cells[0].text == "Rank")
    headers = [c.text for c in ranked.rows[0].cells]
    assert "Priority — ΔV (points)" in headers
    assert any(h.startswith("Indicative cost") for h in headers)
    # And no cell anywhere divides score-points by currency into a single ROI figure.
    text = _all_text(data)
    assert "never divided into one return-on-investment number" in text  # the honesty statement
    assert "ROI" not in text
    assert "/point" not in text and "per point" not in text  # no £-per-ΔV ratio
    assert "£/" not in text


def test_intervention_cost_citing_missing_assumption_refuses() -> None:
    # Traceability is enforced at construction: an intervention cost whose ref is absent from the
    # bridge's register refuses to build the context (§10) — no untraceable currency on the page.
    ctx = _context()
    orphan = RoadmapEntry(
        name="Rogue",
        rank=3,
        delta_v=0.01,
        cost=Money(
            amount_minor=99_000_00, currency=_GBP, assumption_register_ref="AR-DOES-NOT-EXIST"
        ),
    )
    with pytest.raises(ValueError, match="trace to a client-supplied baseline"):
        RoadmapContext(
            subject=ctx.subject,
            bridge=ctx.bridge,
            entries=(*ctx.entries, orphan),
            scenarios=ctx.scenarios,
            engine_version=ctx.engine_version,
            methodology_version=ctx.methodology_version,
            coefficient_version=ctx.coefficient_version,
            uncertainty_version=ctx.uncertainty_version,
            generated_on=ctx.generated_on,
        )


def test_roadmap_is_byte_reproducible() -> None:
    # Same finalised inputs → byte-identical document (the reproducibility the docstrings claim).
    ctx = _context()
    first = build_modernisation_roadmap(ctx, DeliverableMode.DRAFT_INTERNAL)
    second = build_modernisation_roadmap(ctx, DeliverableMode.DRAFT_INTERNAL)
    assert first == second


# ---------------------------------------------------------------- gate + watermark (from GRS-0015)
def test_roadmap_draft_is_watermarked() -> None:
    data = build_modernisation_roadmap(_context(), DeliverableMode.DRAFT_INTERNAL)
    doc = Document(BytesIO(data))
    header = "\n".join(p.text for s in doc.sections for p in s.header.paragraphs)
    assert "DRAFT — not client-usable" in header


def test_service_gate_refuses_client_pack_on_draft_set() -> None:
    ctx = _context()
    coeffs = draft_v1_coefficient_set(_REGISTRY)
    art = compute_score(_doc(graded=True), coeffs, _REGISTRY, _MODEL, random.Random(1))
    with pytest.raises(ClientUsabilityError):
        render_modernisation_roadmap(
            stored_result=art.result,
            coefficients=coeffs,
            bridge=ctx.bridge,
            entries=ctx.entries,
            scenarios=ctx.scenarios,
            uncertainty_version=_MODEL.version,
            subject="Meridian",
            generated_on=date(2026, 7, 13),
            client_facing=True,
        )


def test_service_allows_client_pack_on_client_usable_set() -> None:
    ctx = _context()
    coeffs = _client_usable_set()
    art = compute_score(_doc(graded=True), coeffs, _REGISTRY, _MODEL, random.Random(1))
    rendered = render_modernisation_roadmap(
        stored_result=art.result,
        coefficients=coeffs,
        bridge=ctx.bridge,
        entries=ctx.entries,
        scenarios=ctx.scenarios,
        uncertainty_version=_MODEL.version,
        subject="Meridian",
        generated_on=date(2026, 7, 13),
        client_facing=True,
    )
    assert rendered.mode is DeliverableMode.CLIENT
    assert rendered.docx_bytes[:2] == b"PK"


# ---------------------------------------------------------------- money formatting
def test_format_money_renders_symbol_and_thousands() -> None:
    assert (
        format_money(
            Money(amount_minor=123456700, currency=_GBP, assumption_register_ref="AR-EFFORT")
        )
        == "£1,234,567"
    )
    assert (
        format_money(
            Money(amount_minor=-5000, currency=Currency.USD, assumption_register_ref="AR-EFFORT")
        )
        == "-$50"
    )
    # A sub-pound amount rounds to a clean zero — never a misleading "-£0".
    assert (
        format_money(Money(amount_minor=-30, currency=_GBP, assumption_register_ref="AR-EFFORT"))
        == "£0"
    )
