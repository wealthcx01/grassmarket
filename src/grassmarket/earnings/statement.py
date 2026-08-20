"""Earnings statement export (GRS-0028, PRD §7) — a consultant's own earnings as a .docx, via the
report stack (python-docx). A statement of the caller's OWN lines only; the £ is always shown from
the sealed `Money` figures, never recomputed.
"""

from __future__ import annotations

from collections.abc import Sequence
from datetime import date
from io import BytesIO

from bcap_contracts.commissions import CommissionLine, ConsultancyCommissionCarrot, EarningsSummary
from bcap_contracts.money import Money
from docx import Document


def _fmt(money: Money) -> str:
    """A £-style figure from integer minor units (two decimal places, no locale grouping)."""
    return f"{money.currency.value} {money.amount_minor / 100:,.2f}"


def build_earnings_statement(
    *,
    summary: EarningsSummary,
    lines: Sequence[CommissionLine],
    consultant_name: str,
    generated_on: date,
    consultancy_carrots: Sequence[ConsultancyCommissionCarrot] = (),
) -> bytes:
    """A .docx earnings statement: a header, a per-line table, the summary totals, and the
    Stream-B rate card (GRS-0187) so a statement showing no consultancy still says how consulting
    would be paid. Defaults to empty so an older caller keeps working unchanged."""
    doc = Document()
    doc.add_heading(f"Earnings statement — {consultant_name}", level=0)
    doc.add_paragraph(f"Generated {generated_on.isoformat()}. Amounts in {summary.currency.value}.")

    if lines:
        table = doc.add_table(rows=1, cols=4)
        header = table.rows[0].cells
        header[0].text = "Earned"
        header[1].text = "Kind"
        header[2].text = "Amount"
        header[3].text = "Status"
        for line in lines:
            cells = table.add_row().cells
            cells[0].text = line.earned_on.isoformat() if line.earned_on else "—"
            cells[1].text = line.kind.value.replace("_", " ")
            cells[2].text = _fmt(line.amount)
            cells[3].text = line.payment_status.value
    else:
        doc.add_paragraph("No commission lines recorded.")

    doc.add_heading("Summary", level=1)
    for label, value in (
        ("Earned year-to-date", summary.ytd_earned),
        ("Pending", summary.pending),
        ("Invoiced", summary.invoiced),
        ("Paid", summary.paid),
        ("Projected (earned, unpaid)", summary.projected_unpaid),
    ):
        doc.add_paragraph(f"{label}: {_fmt(value)}")

    # The rate card, not a total: a statement with no consultancy lines should still tell an
    # advisor how consulting pays, which is what the earnings page failed to do (GRS-0187).
    if consultancy_carrots:
        # GRS-0240 scope 2. The page and this document must use one vocabulary: an advisor comparing
        # the two should not have to guess whether "Stream B" here is the same thing as the
        # consulting section there. The page now names both streams; so does this.
        doc.add_heading("Delivering consulting (Stream B)", level=1)
        doc.add_paragraph(
            f"Read live from the {consultancy_carrots[0].schedule_version} schedule. The Year-1 "
            f"rate applies for the first twelve months; the ongoing rate applies after that and "
            f"is uncapped."
        )
        for carrot in consultancy_carrots:
            doc.add_paragraph(
                f"{carrot.delivery_label} · {carrot.sourcing_label}: "
                f"{carrot.yr1_bps / 100:g}% first year, "
                f"{carrot.thereafter_bps / 100:g}% thereafter "
                f"(e.g. {_fmt(carrot.yr1_commission)} then "
                f"{_fmt(carrot.thereafter_commission)} on a "
                f"{_fmt(carrot.example_deal)} engagement)."
            )

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
