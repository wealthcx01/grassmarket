"""Score Evolution Report (GRS-0018, PRD §5). Compares two or more finalised runs of the same
subject: the V/B/P/L trajectory, the deltas between consecutive runs, and — crucially — an explicit
annotation whenever the methodology or coefficient version changed between runs, so a movement is
never silently attributed to the client when it was really a re-weighting (auditability, #6).
"""

from __future__ import annotations

from collections.abc import Sequence
from dataclasses import dataclass
from datetime import date
from io import BytesIO

from bcap_contracts.deliverables import DeliverableMode
from docx.shared import Inches

from grassmarket.atlas.results import AtlasResult
from grassmarket.deliverables.charts import evolution_lines
from grassmarket.deliverables.docx_base import save_document, start_document
from grassmarket.deliverables.uncertainty_text import to_display


@dataclass(frozen=True)
class EvolutionRun:
    """One point in the trajectory: a human label and the immutable stored result it came from."""

    label: str
    result: AtlasResult


def _delta(prev: float, now: float) -> str:
    return f"{(now - prev):+.1f}"


def _version_note(prev: AtlasResult, now: AtlasResult) -> str | None:
    """A one-line annotation when the methodology or coefficient version changed between two runs —
    the reason a delta may be a re-weighting, not a client change."""
    changes = []
    if prev.methodology_version != now.methodology_version:
        changes.append(f"methodology {prev.methodology_version}→{now.methodology_version}")
    if prev.coefficient_version != now.coefficient_version:
        changes.append(f"coefficients {prev.coefficient_version}→{now.coefficient_version}")
    if not changes:
        return None
    return "version change (" + "; ".join(changes) + ") — deltas partly reflect re-weighting"


def build_score_evolution(
    *,
    subject: str,
    runs: Sequence[EvolutionRun],
    mode: DeliverableMode,
    generated_on: date,
) -> bytes:
    """Render the Score Evolution Report. Requires at least two runs (a trajectory needs two
    points); fails loud on a single run rather than draw a line from nothing.

    NOTE (GRS-0021): this multi-run type has no render/service path or endpoint yet. When one is
    wired for a CLIENT-mode Score Evolution pack, it MUST run the §8 committee gate
    (`assert_committee_approved`) first, as the single-run and roadmap render paths do — else a
    client pack could carry an unsigned-off high-stakes rating."""
    if len(runs) < 2:
        raise ValueError("Score Evolution needs at least two finalised runs to compare.")

    doc = start_document(
        title="Score Evolution", subject=subject, generated_on=generated_on, mode=mode
    )

    doc.add_heading("Trajectory", level=1)
    table = doc.add_table(rows=1, cols=6)
    hdr = table.rows[0].cells
    for i, label in enumerate(("Run", "V", "B", "P", "L", "ΔV vs prior")):
        hdr[i].text = label
    prev: AtlasResult | None = None
    for run in runs:
        c = run.result.composite
        cells = table.add_row().cells
        cells[0].text = run.label
        cells[1].text = f"{to_display(c.v_index):.1f}"
        cells[2].text = f"{to_display(c.b_index):.1f}"
        cells[3].text = f"{to_display(c.p_index):.1f}"
        cells[4].text = f"{to_display(c.l_index):.1f}"
        cells[5].text = (
            _delta(to_display(prev.composite.v_index), to_display(c.v_index))
            if prev is not None
            else "—"
        )
        prev = run.result

    doc.add_heading("Version changes between runs", level=1)
    any_change = False
    for earlier, later in zip(runs, runs[1:], strict=False):
        note = _version_note(earlier.result, later.result)
        if note is not None:
            any_change = True
            doc.add_paragraph(f"{earlier.label} → {later.label}: {note}.", style="List Bullet")
    if not any_change:
        doc.add_paragraph(
            "No methodology or coefficient version change between runs — deltas reflect the "
            "subject's own movement.",
            style="List Bullet",
        )

    png = evolution_lines(
        run_labels=[r.label for r in runs],
        series={
            "V": [to_display(r.result.composite.v_index) for r in runs],
            "B": [to_display(r.result.composite.b_index) for r in runs],
            "P": [to_display(r.result.composite.p_index) for r in runs],
            "L": [to_display(r.result.composite.l_index) for r in runs],
        },
    )
    doc.add_picture(BytesIO(png), width=Inches(6.5))
    return save_document(doc)
