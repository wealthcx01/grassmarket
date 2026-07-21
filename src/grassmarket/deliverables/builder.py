"""The python-docx deliverable builder (GRS-0015, PRD §5, SD3 report-stack pattern).

Renders the first sections of the Diagnostic pack from a finalised scoring run:

- **Platform Power Report** — B/P/L/V with honest uncertainty statements (ADR-0008) and the triad
  ordinals with rationale (the ordinal is what a client sees; the audit score is noted, ADR-0002).
- **Methods Appendix** — engine/methodology/coefficient/uncertainty versions, the weight
  elicitation + review dates from the provenance records, and the weight-stability summary.

The document MODE is decided by the client-usable gate (`gate.py`) before we get here: a
DRAFT_INTERNAL document is watermarked "DRAFT — not client-usable" on every page.
"""

from __future__ import annotations

from collections.abc import Sequence
from dataclasses import dataclass
from datetime import date
from io import BytesIO

from bcap_contracts.assessments import CoefficientSet, WidgetObservation
from bcap_contracts.committee import CommitteeDecision, CommitteeDecisionStatus, CommitteeItemType
from bcap_contracts.common import WeightMethod
from bcap_contracts.deliverables import DeliverableMode
from bcap_contracts.narratives import AINarrative, NarrativeStatus
from bcap_contracts.predictions import CBenchmarkRow
from bcap_contracts.registry import WidgetDef
from docx import Document
from docx.document import Document as DocxDocument
from docx.shared import Pt, RGBColor

from grassmarket.atlas.montecarlo import UncertaintyResult
from grassmarket.atlas.results import AtlasResult
from grassmarket.deliverables.gate import DRAFT_WATERMARK
from grassmarket.deliverables.uncertainty_text import format_index_statement, to_display


@dataclass(frozen=True)
class DeliverableContext:
    """Everything the Platform Power Report needs — all derived from the finalised run, so the
    document is reproducible. No `Money` here: the Platform Power Report is score-domain only."""

    subject: str
    result: AtlasResult
    uncertainty: UncertaintyResult
    coefficients: CoefficientSet
    uncertainty_version: str
    generated_on: date
    # Rating Committee calls on this run's high-stakes items (GRS-0021, §8). The approved triad
    # rationale is the text a client sees; every decision + dissent renders into the appendix.
    committee_decisions: tuple[CommitteeDecision, ...] = ()
    # Customer-Proposition (C) inputs (ADR-0023 / GRS-0085), all optional — the C sections render
    # ONLY when `result.customer` is present, and omit cleanly (no blanks/zeros) when it is not.
    # `c_peers` are the APPROVED benchmark rows (GRS-0084); `widgets` are the subject's captured
    # Level-1 grid; `c_widget_defs` supply each widget's rarity/category from the registry.
    c_peers: tuple[CBenchmarkRow, ...] = ()
    widgets: tuple[WidgetObservation, ...] = ()
    c_widget_defs: tuple[WidgetDef, ...] = ()


_TRIAD_ORDER = (
    ("Economic Value", "economic_value"),
    ("Perceived Value", "perceived_value"),
    ("Defence Value", "defence_value"),
)


def build_platform_power_report(
    context: DeliverableContext,
    mode: DeliverableMode,
    narratives: Sequence[AINarrative] = (),
) -> bytes:
    """Render the Platform Power Report + Methods Appendix to .docx bytes. Any approved (or, for an
    internal draft, proposed) AI narratives are rendered into the methods appendix with their
    approval trail (GRS-0017); the client-usable gate upstream guarantees a CLIENT pack only ever
    carries fully-approved narratives."""
    doc = Document()
    if mode is DeliverableMode.DRAFT_INTERNAL:
        _apply_draft_watermark(doc)

    doc.add_heading(f"Platform Power Report — {context.subject}", level=0)
    doc.add_paragraph(f"Generated {context.generated_on.isoformat()}.")

    _platform_power_section(doc, context)
    _customer_proposition_section(doc, context)  # ADR-0023 / GRS-0085 — omits cleanly when no C
    _differentiation_rarity_section(doc, context)
    _methods_appendix(doc, context)
    append_narrative_appendix(doc, narratives, mode)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _apply_draft_watermark(doc: DocxDocument) -> None:
    """A loud DRAFT banner in the page header (every page) plus a red top line. A DRAFT_INTERNAL
    document must never be mistaken for a client pack."""
    header = doc.sections[0].header
    hp = header.paragraphs[0]
    hp.text = DRAFT_WATERMARK
    for run in hp.runs:
        run.bold = True
        run.font.color.rgb = RGBColor(0x8A, 0x20, 0x20)
    banner = doc.add_paragraph()
    run = banner.add_run(DRAFT_WATERMARK)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x8A, 0x20, 0x20)


def _platform_power_section(doc: DocxDocument, context: DeliverableContext) -> None:
    doc.add_heading("Platform Power", level=1)
    unc = context.uncertainty

    doc.add_paragraph(
        "Platform value V and its components, with modelled uncertainty. Where an index's inputs "
        "carried no confidence grade, uncertainty is NOT modelled and the figure is stated as a "
        "point, never a tight range (Methodology §7)."
    )
    comp = context.result.composite
    for name, point, band in (
        ("V (Platform Value)", comp.v_index, unc.v_band),
        ("B (Business)", comp.b_index, unc.b_band),
        ("P (Powers)", comp.p_index, unc.p_band),
        ("L (Platform/Infrastructure)", comp.l_index, unc.l_band),
    ):
        doc.add_paragraph(format_index_statement(name, point, band), style="List Bullet")

    doc.add_paragraph(
        f"Overall assessment uncertainty: {unc.overall_uncertainty.value}.",
    )

    doc.add_heading("Platform Power triad", level=2)
    doc.add_paragraph(
        "The triad is reported as an ORDINAL rating — the words a client sees. The audit-only "
        "score is shown for traceability; ordinal and currency are never mixed (ADR-0002)."
    )
    triad = context.result.triad
    approved_triad = {
        d.item_key: d
        for d in context.committee_decisions
        if d.item_type is CommitteeItemType.TRIAD and d.status is CommitteeDecisionStatus.APPROVED
    }
    for label, attr in _TRIAD_ORDER:
        dim = getattr(triad, attr)
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{label}: {dim.rating}. ").bold = True
        decision = approved_triad.get(attr)
        if decision is not None and decision.rating == dim.rating:
            # The committee-approved rationale is the text a client sees for a high-stakes rating
            # (Methodology §8, scope item 5); the derivation note stays for traceability.
            p.add_run(decision.rationale)
        else:
            p.add_run(
                f"Ordinal rating derived from the continuous inputs against the ADR-0007 triad "
                f"thresholds (audit-only score {dim.score:.3f})."
            )


def _customer_proposition_section(doc: DocxDocument, context: DeliverableContext) -> None:
    """The proposition heatmap (ADR-0023 / GRS-0085): the 10 C modules, the subject's band + q_m,
    and the peer-benchmark average per module. Renders ONLY when the run carries a C result; a run
    with no C is omitted cleanly — never a blank table or a zero-filled row (Stage 1: C is reported
    alongside V, not part of the headline V)."""
    customer = context.result.customer
    if customer is None:
        return

    doc.add_heading("Customer Proposition (C)", level=1)
    doc.add_paragraph(
        f"The Customer-Proposition index C is {to_display(customer.value):.1f} (0–100). C is "
        "reported ALONGSIDE Platform Value V (ADR-0023); it is not part of the headline V in this "
        "release. The heatmap below reads the ten proposition modules against the peer set."
    )

    # Peer average per C module (only over peers that scored it) — first-class absence, so a
    # module no peer scored simply shows "—", never a fabricated 0.
    peer_means: dict[str, float] = {}
    for module in customer.modules:
        peer_values = [
            p.module_scores[module.key] for p in context.c_peers if module.key in p.module_scores
        ]
        if peer_values:
            peer_means[module.key] = sum(peer_values) / len(peer_values)

    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = ("Module", "Rating", "Score", "Peer avg")
    for module in customer.modules:
        row = table.add_row().cells
        row[0].text = module.name
        row[1].text = module.gate_band
        row[2].text = f"{to_display(module.q_m):.1f}" if module.q_m is not None else "Not Assessed"
        row[3].text = (
            f"{to_display(peer_means[module.key]):.1f}" if module.key in peer_means else "—"
        )

    if context.c_peers:
        subject_c = customer.value
        ahead = sum(1 for p in context.c_peers if p.c_index < subject_c)
        doc.add_paragraph(
            f"Against {len(context.c_peers)} benchmarked peer(s), the subject's proposition ranks "
            f"ahead of {ahead}."
        )


def _differentiation_rarity_section(doc: DocxDocument, context: DeliverableContext) -> None:
    """The differentiation-vs-rarity map (ADR-0023 / GRS-0085): widgets read by RARITY against
    presence. A Rare widget the subject HAS is a differentiation asset; a Common one it LACKS is a
    table-stakes gap. Renders only when both a C result and captured widgets exist."""
    if context.result.customer is None or not context.widgets or not context.c_widget_defs:
        return

    rarity_by_key = {w.key: w.rarity for w in context.c_widget_defs}
    name_by_key = {w.key: w.name for w in context.c_widget_defs}
    present = {o.widget_key for o in context.widgets if o.present}
    # A captured-but-not-present widget is a gap unless explicitly paywalled/defective, which we
    # surface separately (present-yet-not-a-clean-pass).
    flagged = {
        o.widget_key: o.state for o in context.widgets if not o.present and o.state is not None
    }
    absent = {o.widget_key for o in context.widgets if not o.present}

    differentiators = sorted(
        k for k in present if rarity_by_key.get(k) == "Rare" and k in name_by_key
    )
    table_stakes_gaps = sorted(
        k for k in absent if rarity_by_key.get(k) == "Common" and k in name_by_key
    )

    doc.add_heading("Differentiation vs. rarity", level=1)
    doc.add_paragraph(
        "The Level-1 widget checklist read against rarity. A RARE widget the subject offers is a "
        "differentiation asset; a COMMON widget it lacks is a table-stakes gap. Uncommon ones and "
        "present-but-paywalled/defective features are noted for context."
    )

    doc.add_heading("Differentiation assets (Rare, present)", level=2)
    if differentiators:
        for k in differentiators:
            doc.add_paragraph(name_by_key[k], style="List Bullet")
    else:
        doc.add_paragraph("No Rare widget is offered — no rarity-driven differentiation captured.")

    doc.add_heading("Table-stakes gaps (Common, absent)", level=2)
    if table_stakes_gaps:
        for k in table_stakes_gaps:
            doc.add_paragraph(name_by_key[k], style="List Bullet")
    else:
        doc.add_paragraph("No Common widget is missing — table-stakes coverage is complete.")

    paywalled_or_defective = sorted(k for k in flagged if k in name_by_key)
    if paywalled_or_defective:
        doc.add_heading("Present but gated / defective", level=2)
        for k in paywalled_or_defective:
            doc.add_paragraph(f"{name_by_key[k]} — {flagged[k].value}", style="List Bullet")


def _methods_appendix(doc: DocxDocument, context: DeliverableContext) -> None:
    doc.add_heading("Methods Appendix", level=1)
    result = context.result

    doc.add_heading("Versions", level=2)
    for label, value in (
        ("Engine", result.engine_version),
        ("Methodology", result.methodology_version),
        ("Coefficient set", result.coefficient_version),
        ("Uncertainty model", context.uncertainty_version),
    ):
        doc.add_paragraph(f"{label}: {value}", style="List Bullet")

    doc.add_heading("Weight provenance", level=2)
    provenance = context.coefficients.provenance
    theta = provenance.get("theta")
    if theta is not None:
        if theta.method is WeightMethod.DIRECT:
            # A DIRECT weight family is a draft placeholder, not panel output — say so plainly
            # rather than claim elicitation. (Draft sets only render watermarked internal docs.)
            doc.add_paragraph(
                f"Draft placeholder weights (method {theta.method.value}, not expert-elicited), "
                f"set {theta.set_on.isoformat()}, review due {theta.review_due.isoformat()}."
            )
        else:
            # Attribute the method to θ specifically — families differ (θ/α swing, λ/δ AHP, strength
            # Delphi); the per-family table below carries each family's own record.
            doc.add_paragraph(
                f"Weights expert-elicited {theta.set_on.isoformat()} (headline θ by "
                f"{theta.method.value}), review due {theta.review_due.isoformat()}. "
                f"Per-family method and dispersion below."
            )

    doc.add_heading("Weight-stability summary", level=2)
    doc.add_paragraph(
        "Recorded dispersion for each coefficient family — the spread the §7 stability interval is "
        "drawn from."
    )
    table = doc.add_table(rows=1, cols=5)
    hdr = table.rows[0].cells
    hdr[0].text = "Family"
    hdr[1].text = "Method"
    hdr[2].text = "Set on"
    hdr[3].text = "Dispersion"
    hdr[4].text = "Review due"
    for family in sorted(provenance):
        record = provenance[family]
        row = table.add_row().cells
        row[0].text = family
        row[1].text = record.method.value
        row[2].text = record.set_on.isoformat()
        row[3].text = record.dispersion
        row[4].text = record.review_due.isoformat()

    doc.add_paragraph(
        f"Headline platform value V = {to_display(result.composite.v_index):.1f} "
        f"(display scale 0–100)."
    )

    if result.customer is not None:
        # C provenance note (ADR-0023 / GRS-0085): make the staged, reported-not-summed status
        # explicit, and that peer benchmarks are approval-gated (ADR-0009), so no reader mistakes C
        # for part of V or the peer set for auto-ingested data.
        doc.add_heading("Customer Proposition (C) — provenance", level=2)
        doc.add_paragraph(
            f"C = {to_display(result.customer.value):.1f} is computed on the same rubric family as "
            "the infrastructure index over the separate Phase-E module set (ADR-0023). In this "
            "release C is REPORTED ALONGSIDE V and is NOT part of the headline V. The C "
            "coefficients are draft (a θ_C elicitation panel is post-launch). Peer benchmarks are "
            "approval-gated (ADR-0009): only consultant-approved peer rows appear here."
        )

    _committee_appendix(doc, context)


def _committee_appendix(doc: DocxDocument, context: DeliverableContext) -> None:
    """Render the Rating Committee's calls on this run's high-stakes ratings — recorded rationale
    and any dissent (Methodology §8 audit evidence). Absent for an assessment with no high-stakes
    ratings (nothing needed sign-off)."""
    if not context.committee_decisions:
        return
    doc.add_heading("Rating Committee decisions", level=2)
    doc.add_paragraph(
        "High-stakes ratings (a power Established or above, a triad rating above None, a module "
        "rated Frontier) carry recorded committee sign-off with rationale and dissent — judgment "
        "disciplined by peer challenge, not formula (Methodology §8)."
    )
    for decision in context.committee_decisions:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(
            f"{decision.item_type.value.title()} '{decision.item_key}' ({decision.rating}) — "
            f"{decision.status.value}: "
        ).bold = True
        p.add_run(decision.rationale)
        trail = doc.add_paragraph()
        trail.add_run(
            f"    Decided by committee member {decision.decided_by_consultant_id} at "
            f"{decision.decided_at.isoformat()}."
        )
        if decision.dissent_note is not None:
            dissent = doc.add_paragraph()
            dissent.add_run(f"    Dissent: {decision.dissent_note}").italic = True


AI_DRAFTED_LABEL = "AI-DRAFTED"


def append_narrative_appendix(
    doc: DocxDocument, narratives: Sequence[AINarrative], mode: DeliverableMode
) -> None:
    """Render the AI-drafted narrative sections and their approval trail into the methods appendix
    (GRS-0017, §5). Every AI section is labelled ``AI-DRAFTED``; an approved section prints who
    approved it, when, and the edit summary; an unapproved section is flagged as not client-usable.

    A CLIENT-mode document only ever reaches here with fully-approved narratives (the gate refuses
    otherwise), so the approval line is always present for a client pack."""
    if not narratives:
        return
    doc.add_heading("AI-drafted narratives", level=2)
    doc.add_paragraph(
        "The interpretive sections below were AI-drafted and are gated: nothing reaches a client "
        "without a recorded human approval (non-negotiable #8)."
    )
    for narrative in narratives:
        heading = doc.add_paragraph(style="List Bullet")
        heading.add_run(f"[{AI_DRAFTED_LABEL}] {narrative.section.value.title()}: ").bold = True
        if narrative.status is NarrativeStatus.APPROVED:
            # The contract guarantees an APPROVED narrative carries its full trail — assert it,
            # never silently default a blank/"unknown" (non-negotiable #3: fail loud).
            assert narrative.final_text is not None and narrative.approved_at is not None
            heading.add_run(narrative.final_text)
            doc.add_paragraph(
                f"    Approved by consultant {narrative.approved_by_consultant_id} at "
                f"{narrative.approved_at.isoformat()}. Edits: {narrative.edit_summary}."
            )
        else:
            heading.add_run(narrative.proposed_text)
            doc.add_paragraph(
                f"    Status: {narrative.status.value} — not client-usable until approved."
            )
