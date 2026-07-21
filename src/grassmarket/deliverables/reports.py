"""The narrative Diagnostic-pack documents that derive from a single finalised run (GRS-0018):
Executive Summary, Technical Appendix, and the Workshop Output template. Each is score-domain only
(no currency — that is the value bridge, GRS-0016) and honest about uncertainty (Methodology §7) and
Not Assessed coverage (§3.2). All render deterministically from the immutable stored result.
"""

from __future__ import annotations

from collections.abc import Sequence
from io import BytesIO

from bcap_contracts.common import MaturityLevel, NonScoreState
from bcap_contracts.deliverables import DeliverableMode
from bcap_contracts.narratives import AINarrative
from docx.document import Document as DocxDocument
from docx.shared import Inches

from grassmarket.atlas.montecarlo import Band
from grassmarket.deliverables.builder import DeliverableContext
from grassmarket.deliverables.charts import index_tornado
from grassmarket.deliverables.docx_base import save_document, start_document
from grassmarket.deliverables.uncertainty_text import format_index_statement, to_display


def _start(context: DeliverableContext, mode: DeliverableMode, title: str) -> DocxDocument:
    return start_document(
        title=title, subject=context.subject, generated_on=context.generated_on, mode=mode
    )


def _finish(doc: DocxDocument) -> bytes:
    return save_document(doc)


def _tornado_png(context: DeliverableContext) -> bytes:
    unc = context.uncertainty
    named: list[tuple[str, Band]] = [
        ("V", unc.v_band),
        ("B", unc.b_band),
        ("P", unc.p_band),
        ("L", unc.l_band),
    ]
    # An unmodelled index collapses to a point (low==mid==high) — honest, never a false band.
    lows = [to_display(b.p10 if b.modelled else b.p50) for _, b in named]
    mids = [to_display(b.p50) for _, b in named]
    highs = [to_display(b.p90 if b.modelled else b.p50) for _, b in named]
    return index_tornado(labels=[name for name, _ in named], lows=lows, mids=mids, highs=highs)


# --------------------------------------------------------------------- Executive Summary
def build_executive_summary(
    context: DeliverableContext, mode: DeliverableMode, narratives: Sequence[AINarrative] = ()
) -> bytes:
    """A short, board-ready summary: the headline, the triad, strengths, risks, uncertainty. AI
    narratives are not rendered here (they belong to the Platform Power Report); the parameter
    keeps the single-run builder signature uniform for the dispatcher."""
    del narratives
    doc = _start(context, mode, "Executive Summary")
    result = context.result
    unc = context.uncertainty

    doc.add_heading("Headline", level=1)
    doc.add_paragraph(
        format_index_statement("Platform value V", result.composite.v_index, unc.v_band)
    )
    doc.add_paragraph(f"Overall assessment uncertainty: {unc.overall_uncertainty.value}.")

    doc.add_heading("Platform Power triad", level=1)
    triad = result.triad
    for label, dim in (
        ("Economic Value", triad.economic_value),
        ("Perceived Value", triad.perceived_value),
        ("Defence Value", triad.defence_value),
    ):
        p = doc.add_paragraph(style="List Bullet")
        word = dim.rating if dim.rating is not None else "Not assessed"
        p.add_run(f"{label}: {word}.").bold = True

    doc.add_heading("Where the platform is strong", level=1)
    _STRONG = (MaturityLevel.ADVANCED.value, MaturityLevel.FRONTIER.value)
    strong = [m for m in result.modules if m.gate_band in _STRONG]
    if strong:
        for m in strong:
            doc.add_paragraph(f"{m.name}: {m.gate_band}.", style="List Bullet")
    else:
        doc.add_paragraph("No module reaches Advanced yet.", style="List Bullet")

    doc.add_heading("Where the risk sits", level=1)
    risks = [
        m for m in result.modules if m.gate_blocked or m.gate_band == MaturityLevel.BASIC.value
    ]
    if risks:
        for m in risks:
            reason = (
                "gate blocked (a critical subcomponent is Not Assessed)"
                if m.gate_blocked
                else ("rated Basic")
            )
            doc.add_paragraph(f"{m.name}: {reason}.", style="List Bullet")
    else:
        doc.add_paragraph("No module is gate-blocked or Basic.", style="List Bullet")

    doc.add_heading("Uncertainty at a glance", level=1)
    doc.add_paragraph(
        "P10–P90 ranges per index (P50 marked). An index whose inputs carried no confidence grade "
        "is shown as a point, never a false-tight band (Methodology §7)."
    )
    doc.add_picture(BytesIO(_tornado_png(context)), width=Inches(6.0))
    return _finish(doc)


# --------------------------------------------------------------------- Technical Appendix
def build_technical_appendix(
    context: DeliverableContext, mode: DeliverableMode, narratives: Sequence[AINarrative] = ()
) -> bytes:
    """The full auditable detail: versions, weight provenance, per-module gate reasoning, and the
    L / B / P composition. Extends the GRS-0015 methods appendix (PRD §5). (``narratives`` keeps the
    single-run builder signature uniform for the dispatcher; not rendered here.)"""
    del narratives
    doc = _start(context, mode, "Technical Appendix")
    result = context.result

    doc.add_heading("Versions", level=1)
    for label, value in (
        ("Engine", result.engine_version),
        ("Methodology", result.methodology_version),
        ("Coefficient set", result.coefficient_version),
        ("Uncertainty model", context.uncertainty_version),
    ):
        doc.add_paragraph(f"{label}: {value}", style="List Bullet")

    doc.add_heading("Weight provenance", level=1)
    provenance = context.coefficients.provenance
    for family in sorted(provenance):
        record = provenance[family]
        doc.add_paragraph(
            f"{family}: elicited {record.set_on.isoformat()} ({record.method.value}), "
            f"dispersion {record.dispersion}, review due {record.review_due.isoformat()}.",
            style="List Bullet",
        )

    doc.add_heading("Module gate reasoning", level=1)
    table = doc.add_table(rows=1, cols=5)
    hdr = table.rows[0].cells
    for i, label in enumerate(("Module", "Gate", "q_m (0–100)", "Coverage", "Note")):
        hdr[i].text = label
    for m in result.modules:
        cells = table.add_row().cells
        cells[0].text = m.name
        cells[1].text = m.gate_band + (" (blocked)" if m.gate_blocked else "")
        cells[2].text = f"{to_display(m.q_m):.1f}" if m.q_m is not None else "n/a"
        cells[3].text = f"{m.n_assessed}/{m.n_applicable}"
        cells[4].text = m.gate_note

    doc.add_heading("Composition", level=1)
    lr = result.l_index
    doc.add_paragraph(
        f"L = {to_display(lr.value):.1f} (weighted term {to_display(lr.weighted_term):.1f}, "
        f"min term {to_display(lr.min_term):.1f}).",
        style="List Bullet",
    )
    doc.add_paragraph(
        f"B = {to_display(result.business.b_index):.1f}; "
        f"P = {to_display(result.powers.p_index):.1f}.",
        style="List Bullet",
    )
    return _finish(doc)


# --------------------------------------------------------------------- Workshop Output
def build_workshop_output(
    context: DeliverableContext, mode: DeliverableMode, narratives: Sequence[AINarrative] = ()
) -> bytes:
    """A pre-engagement workshop template. Tolerates partial data: it reports what is assessed and
    what is still Not Assessed, shows the (wide) uncertainty, and leaves discussion prompts. It
    never pretends coverage it does not have (§3.2). (``narratives`` keeps the builder signature
    uniform for the dispatcher; not rendered here.)"""
    del narratives
    doc = _start(context, mode, "Workshop Output")
    result = context.result

    doc.add_paragraph(
        "Pre-engagement mode: this template is designed for partial data. Unassessed subcomponents "
        "are listed as discussion inputs, not gaps to hide, and uncertainty is expected to be wide."
    )

    doc.add_heading("Coverage", level=1)
    for m in result.modules:
        not_assessed = [
            r.key for r in m.subcomponents if r.state == NonScoreState.NOT_ASSESSED.value
        ]
        line = f"{m.name}: {m.n_assessed}/{m.n_applicable} assessed"
        if not_assessed:
            line += f" — to discuss: {', '.join(not_assessed)}"
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("Current reading (wide, provisional)", level=1)
    comp = context.result.composite
    for name, point, band in (
        ("V", comp.v_index, context.uncertainty.v_band),
        ("B", comp.b_index, context.uncertainty.b_band),
        ("P", comp.p_index, context.uncertainty.p_band),
        ("L", comp.l_index, context.uncertainty.l_band),
    ):
        doc.add_paragraph(format_index_statement(name, point, band), style="List Bullet")

    doc.add_heading("Workshop discussion prompts", level=1)
    for prompt in (
        "Which Not Assessed subcomponents can we evidence during the engagement?",
        "Where does the client dispute the provisional reading, and on what evidence?",
        "Which gate-blocking criticals are the priority to close first?",
    ):
        doc.add_paragraph(prompt, style="List Bullet")
    return _finish(doc)
