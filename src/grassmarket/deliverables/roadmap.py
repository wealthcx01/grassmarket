"""The Modernisation Roadmap — the flagship money pages of the Diagnostic pack (GRS-0016, §10).

Honest by construction (ADR-0002): the Roadmap RANKS interventions by the Upgrade Priority Index
(ΔV, score domain) and PRICES them with the value bridge (currency domain), and the two are shown
side by side — never divided into a single ROI number. The prototype's `LV = κ·Δq/(1+r) − cost`
(defect D2) is not expressible: no function here takes both a `Score` and a `Money`, the ΔV arrives
as a plain float, and every currency figure resolves to a printed assumption-register row.

Rendering only. The domain objects — the `ValueBridge`, the priority index, the scenario ΔVs — are
computed upstream (`grassmarket.value`) from a finalised run, so the document is reproducible.
"""

from __future__ import annotations

from dataclasses import dataclass
from datetime import date
from io import BytesIO

from bcap_contracts.deliverables import DeliverableMode
from bcap_contracts.money import Money
from bcap_contracts.value import LeverKind, ScenarioResult, ValueBridge
from docx import Document
from docx.document import Document as DocxDocument
from docx.shared import Inches

from grassmarket.deliverables.builder import _apply_draft_watermark
from grassmarket.deliverables.charts import priority_cost_scatter
from grassmarket.deliverables.money_text import currency_symbol, format_money, major_units
from grassmarket.deliverables.uncertainty_text import to_display

_LEVER_LABEL = {
    LeverKind.COST_TO_SERVE: "Cost to serve",
    LeverKind.PROJECT_DRAG: "Project drag recovered",
    LeverKind.INCIDENT_EXPECTED_LOSS: "Incident expected loss avoided",
    LeverKind.REVENUE_ENABLEMENT: "Revenue enablement",
}


@dataclass(frozen=True)
class RoadmapEntry:
    """One ranked intervention: its priority (ΔV, score domain) and its indicative cost (currency).
    The two live in one record for side-by-side display — they are never combined into a ratio.

    The cost is a `Money`, so it carries its own assumption-register ref: an indicative cost is a
    currency figure like any other and must trace to a printed baseline (§10). It is never a bare
    integer that appears on the page from nowhere."""

    name: str
    rank: int
    delta_v: float  # signed difference of two [0,1] scores — score domain, NOT a Score
    cost: Money  # register-linked indicative cost, for the ranked table + priority-vs-cost scatter


@dataclass(frozen=True)
class RoadmapContext:
    """Everything the Roadmap needs, all derived from a finalised run so the doc is reproducible.

    Traceability is enforced at construction (§10, mirroring the bridge's own figure check): every
    intervention cost must cite an assumption in the bridge's register, or the context refuses to
    build — no currency figure reaches the page without a printed baseline."""

    subject: str
    bridge: ValueBridge
    entries: tuple[RoadmapEntry, ...]
    scenarios: tuple[ScenarioResult, ...]
    engine_version: str
    methodology_version: str
    coefficient_version: str
    uncertainty_version: str
    generated_on: date

    def __post_init__(self) -> None:
        legal = self.bridge.assumption_register.refs()
        missing = sorted(
            e.cost.assumption_register_ref
            for e in self.entries
            if e.cost.assumption_register_ref not in legal
        )
        if missing:
            raise ValueError(
                f"Roadmap intervention costs cite assumptions not in the register: {missing}. "
                f"Every currency figure must trace to a client-supplied baseline (Methodology §10)."
            )


def _delta_points(delta_v: float) -> str:
    """ΔV on the 0–100 display scale, always signed."""
    return f"{to_display(delta_v):+.1f}"


def build_modernisation_roadmap(context: RoadmapContext, mode: DeliverableMode) -> bytes:
    """Render the Modernisation Roadmap to .docx bytes. Watermarked when DRAFT_INTERNAL."""
    doc = Document()
    if mode is DeliverableMode.DRAFT_INTERNAL:
        _apply_draft_watermark(doc)

    doc.add_heading(f"Modernisation Roadmap — {context.subject}", level=0)
    doc.add_paragraph(f"Generated {context.generated_on.isoformat()}.")
    doc.add_paragraph(
        "This roadmap RANKS interventions by the Upgrade Priority Index (the change in platform "
        "value ΔV from full re-scoring — the score domain) and PRICES them through the value "
        "bridge (the currency domain). Priority and price are shown side by side and are never "
        "divided into one return-on-investment number: a score-point and a pound do not belong in "
        "the same equation (ADR-0002)."
    )

    _roadmap_section(doc, context)
    _value_bridge_section(doc, context)
    _scatter_section(doc, context)
    _scenario_comparison_section(doc, context)
    _assumption_register_section(doc, context)
    _versions_section(doc, context)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _roadmap_section(doc: DocxDocument, context: RoadmapContext) -> None:
    doc.add_heading("Interventions, ranked", level=1)
    doc.add_paragraph(
        "Ranked by the Upgrade Priority Index (ΔV, descending). The indicative cost sits alongside "
        "each rank from the value bridge — for comparison, not as a ratio."
    )
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "Rank"
    hdr[1].text = "Intervention"
    hdr[2].text = "Priority — ΔV (points)"
    hdr[3].text = "Indicative cost"
    for entry in sorted(context.entries, key=lambda e: e.rank):
        row = table.add_row().cells
        row[0].text = str(entry.rank)
        row[1].text = entry.name
        row[2].text = _delta_points(entry.delta_v)  # score domain
        row[3].text = format_money(entry.cost)  # currency domain — separate column, never a ratio


def _value_bridge_section(doc: DocxDocument, context: RoadmapContext) -> None:
    bridge = context.bridge
    doc.add_heading("Value bridge", level=1)

    doc.add_heading("Layer 1 — Cost", level=2)
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(f"Indicative upgrade cost: {format_money(bridge.cost.total)}. ").bold = True
    if bridge.cost.note:
        p.add_run(bridge.cost.note + " ")
    p.add_run(f"Assumptions: {', '.join(bridge.cost.assumption_refs)}.")

    doc.add_heading("Layer 2 — Levers (risk-adjusted NPV)", level=2)
    doc.add_paragraph(
        "Each evidenced lever's NPV, with the assumption-register refs it traces to. Every figure "
        "resolves to a printed baseline in the register below."
    )
    table = doc.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text = "Lever"
    hdr[1].text = "Risk-adjusted NPV"
    hdr[2].text = "Assumptions"
    for lever in bridge.levers:
        row = table.add_row().cells
        row[0].text = _LEVER_LABEL[lever.lever]
        row[1].text = format_money(lever.npv)
        row[2].text = ", ".join(lever.assumption_refs)
    if bridge.levers:
        # total_lever_npv() refuses an empty tuple (Money can't sum from nothing); the contract
        # permits a lever-less bridge, so guard rather than crash mid-render.
        doc.add_paragraph(
            f"Total risk-adjusted lever NPV: {format_money(bridge.total_lever_npv())} "
            f"(Money summed with Money only — the cost is a separate figure, never netted here)."
        )
    else:
        doc.add_paragraph("No evidenced levers recorded for this subject.")

    doc.add_heading("Layer 3 — Strategic (ordinal)", level=2)
    doc.add_paragraph(
        "Moat and durability implications as ORDINAL ratings, in duration language — never priced "
        "in currency (ADR-0002)."
    )
    for rating in bridge.strategic:
        para = doc.add_paragraph(style="List Bullet")
        para.add_run(f"{rating.dimension}: {rating.rating}. ").bold = True
        para.add_run(rating.rationale)


def _scatter_section(doc: DocxDocument, context: RoadmapContext) -> None:
    doc.add_heading("Priority vs cost", level=1)
    doc.add_paragraph(
        "Each intervention plotted against its indicative cost (horizontal, currency) and its "
        "Upgrade Priority Index ΔV (vertical, score points). Two axes, never one ratio."
    )
    ordered = sorted(context.entries, key=lambda e: e.rank)
    png = priority_cost_scatter(
        labels=[e.name for e in ordered],
        costs_major=[major_units(e.cost) for e in ordered],
        priority_points=[to_display(e.delta_v) for e in ordered],
        currency_symbol=currency_symbol(context.bridge.cost.total.currency),
    )
    doc.add_picture(BytesIO(png), width=Inches(6.0))


def _scenario_comparison_section(doc: DocxDocument, context: RoadmapContext) -> None:
    doc.add_heading("Scenario comparison", level=1)
    doc.add_paragraph(
        "Named scenarios evaluated by full re-scoring against the baseline. ΔV is the headline; "
        "the component deltas show where the value moves (all score domain, on the 0–100 scale)."
    )
    table = doc.add_table(rows=1, cols=5)
    hdr = table.rows[0].cells
    for i, label in enumerate(("Scenario", "ΔV", "ΔB", "ΔP", "ΔL")):
        hdr[i].text = label
    for scenario in context.scenarios:
        row = table.add_row().cells
        row[0].text = scenario.name
        row[1].text = _delta_points(scenario.delta_v)
        row[2].text = _delta_points(scenario.delta_b)
        row[3].text = _delta_points(scenario.delta_p)
        row[4].text = _delta_points(scenario.delta_l)


def _assumption_register_section(doc: DocxDocument, context: RoadmapContext) -> None:
    doc.add_heading("Assumption register", level=1)
    doc.add_paragraph(
        "Every currency figure above traces to a client-supplied baseline here (Methodology §10). "
        "No figure appears without its register entry."
    )
    table = doc.add_table(rows=1, cols=5)
    hdr = table.rows[0].cells
    for i, label in enumerate(("Ref", "Baseline", "Unit", "Source", "Description")):
        hdr[i].text = label
    for a in context.bridge.assumption_register.entries:
        row = table.add_row().cells
        row[0].text = a.ref
        row[1].text = f"{a.baseline_value:g}"
        row[2].text = a.unit
        row[3].text = a.source
        row[4].text = a.description


def _versions_section(doc: DocxDocument, context: RoadmapContext) -> None:
    doc.add_heading("Methods & versions", level=1)
    for label, value in (
        ("Engine", context.engine_version),
        ("Methodology", context.methodology_version),
        ("Coefficient set", context.coefficient_version),
        ("Uncertainty model", context.uncertainty_version),
    ):
        doc.add_paragraph(f"{label}: {value}", style="List Bullet")
