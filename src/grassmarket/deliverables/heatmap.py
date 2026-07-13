"""Infrastructure Heatmap (GRS-0018, PRD §5, Methodology §3.2).

A module × subcomponent grid, each rating cell colour-banded by its maturity. The one hard rule:
**Not Assessed is rendered visually distinct — never red, never the same fill as Basic** (§3.2:
a non-score state is first-class, never conflated with a low score). Not Applicable is a third,
lighter fill. The fills are exposed as constants so the rendering test can assert on the doc XML.
"""

from __future__ import annotations

from collections.abc import Sequence
from io import BytesIO

from bcap_contracts.common import MaturityLevel, NonScoreState
from bcap_contracts.deliverables import DeliverableMode
from bcap_contracts.narratives import AINarrative
from docx.document import Document as DocxDocument
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.table import _Cell

from grassmarket.deliverables.builder import DeliverableContext
from grassmarket.deliverables.charts import module_radar
from grassmarket.deliverables.docx_base import save_document, start_document
from grassmarket.deliverables.uncertainty_text import to_display

# Ordered warm→cool by maturity. Basic is a distinct orange — NOT red — so Not Assessed (grey) can
# never be mistaken for it, and neither reads as an alarm colour the palette doesn't otherwise use.
_BAND_FILL: dict[str, str] = {
    MaturityLevel.BASIC.value: "C55A11",  # dark orange
    MaturityLevel.DEVELOPING.value: "FFC000",  # amber
    MaturityLevel.ADVANCED.value: "70AD47",  # green
    MaturityLevel.FRONTIER.value: "2E75B6",  # blue
}
NOT_ASSESSED_FILL = "808080"  # neutral grey — first-class, distinct, never red, never Basic
NOT_APPLICABLE_FILL = "D9D9D9"  # light grey — out of scope
_DEFAULT_FILL = "FFFFFF"


def _rating_fill(row) -> tuple[str, str]:
    """(display text, fill hex) for a subcomponent's rating cell."""
    if row.state == NonScoreState.NOT_ASSESSED.value:
        return "Not Assessed", NOT_ASSESSED_FILL
    if row.state == NonScoreState.NOT_APPLICABLE.value:
        return "Not Applicable", NOT_APPLICABLE_FILL
    if row.level is not None:
        # Fail loud on an unexpected maturity label rather than silently shading it a neutral
        # colour (#3): the contract guarantees one of the four MaturityLevel labels here.
        return row.level, _BAND_FILL[row.level]
    # An assessed/non-score row always has a level or a state; this only covers a genuinely empty
    # row, and prints a neutral placeholder (a display choice, not a silent score default).
    return "—", _DEFAULT_FILL


def _shade_cell(cell: _Cell, fill_hex: str) -> None:
    """Apply a solid background fill to a table cell (docx w:shd) — no python-docx API covers it."""
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    cell._tc.get_or_add_tcPr().append(shd)


def cell_fill(cell: _Cell) -> str | None:
    """Read back a cell's shading fill (used by the rendering test)."""
    tcPr = cell._tc.tcPr
    if tcPr is None:
        return None
    shd = tcPr.find(qn("w:shd"))
    return None if shd is None else shd.get(qn("w:fill"))


def build_infrastructure_heatmap(
    context: DeliverableContext, mode: DeliverableMode, narratives: Sequence[AINarrative] = ()
) -> bytes:
    """Render the Infrastructure Heatmap to .docx bytes. (``narratives`` keeps the single-run
    builder signature uniform for the dispatcher; not rendered here.)"""
    del narratives
    doc = start_document(
        title="Infrastructure Heatmap",
        subject=context.subject,
        generated_on=context.generated_on,
        mode=mode,
    )
    _legend(doc)
    _radar(doc, context)
    _grid(doc, context)
    return save_document(doc)


def _legend(doc: DocxDocument) -> None:
    doc.add_paragraph(
        "Each subcomponent is banded by maturity. Not Assessed is shown in a distinct neutral grey "
        "— it is a first-class state, never scored as zero and never conflated with Basic "
        "(Methodology §3.2). Not Applicable is a lighter grey (out of scope)."
    )


def _radar(doc: DocxDocument, context: DeliverableContext) -> None:
    scoreable = [(m.name, m.q_m) for m in context.result.modules if m.q_m is not None]
    if len(scoreable) < 3:
        return  # a radar needs ≥3 axes; a sparse assessment simply omits it
    png = module_radar(
        labels=[name for name, _ in scoreable],
        values=[to_display(q) for _, q in scoreable],
    )
    doc.add_picture(BytesIO(png), width=Inches(5.5))


def _grid(doc: DocxDocument, context: DeliverableContext) -> None:
    for module in context.result.modules:
        blocked = " — GATE BLOCKED" if module.gate_blocked else ""
        doc.add_heading(f"{module.name} — Gate: {module.gate_band}{blocked}", level=2)
        table = doc.add_table(rows=1, cols=3)
        hdr = table.rows[0].cells
        hdr[0].text = "Subcomponent"
        hdr[1].text = "Rating"
        hdr[2].text = "Critical"
        for row in module.subcomponents:
            cells = table.add_row().cells
            cells[0].text = row.key
            text, fill = _rating_fill(row)
            cells[1].text = text
            _shade_cell(cells[1], fill)
            cells[2].text = "critical" if row.critical else ""
