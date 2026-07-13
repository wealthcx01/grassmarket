"""Shared docx scaffolding for the Diagnostic-pack builders (GRS-0018).

Every single-run document opens the same way — a fresh Document, a DRAFT watermark when the mode is
internal, a title heading, and a "Generated {date}" line — and closes the same way (save to bytes).
Consolidated here so the header/watermark convention lives in ONE place rather than being re-spelled
in each builder. (The GRS-0015 Platform Power Report predates this and still inlines its own copy.)
"""

from __future__ import annotations

from datetime import date
from io import BytesIO

from bcap_contracts.deliverables import DeliverableMode
from docx import Document
from docx.document import Document as DocxDocument

from grassmarket.deliverables.builder import _apply_draft_watermark


def start_document(
    *, title: str, subject: str, generated_on: date, mode: DeliverableMode
) -> DocxDocument:
    """A new document with the standard header: DRAFT watermark (internal mode only), a level-0
    title heading, and the generated-on line."""
    doc = Document()
    if mode is DeliverableMode.DRAFT_INTERNAL:
        _apply_draft_watermark(doc)
    doc.add_heading(f"{title} — {subject}", level=0)
    doc.add_paragraph(f"Generated {generated_on.isoformat()}.")
    return doc


def save_document(doc: DocxDocument) -> bytes:
    """Serialise a document to .docx bytes."""
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
